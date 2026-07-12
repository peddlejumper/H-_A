#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_hsharp_algorithm_examples():
    doc = Document('/Users/peddlejumper/H#/v0.4/资料/HSharp_Programming_Tutorial.docx')

    # A.8 排序算法 / Sorting Algorithms
    p = doc.add_paragraph()
    run = p.add_run('A.8 排序算法 / Sorting Algorithms')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(26, 86, 142)

    # A.8.1 冒泡排序 / Bubble Sort
    p = doc.add_paragraph()
    run = p.add_run('A.8.1 冒泡排序 / Bubble Sort')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('# 冒泡排序实现 / Bubble Sort Implementation\n')
    p.add_run('fn bubble_sort(arr) {\n')
    p.add_run('    let n = len(arr);\n')
    p.add_run('    let i = 0;\n')
    p.add_run('    while (i < n - 1) {\n')
    p.add_run('        let j = 0;\n')
    p.add_run('        while (j < n - i - 1) {\n')
    p.add_run('            if (arr[j] > arr[j + 1]) {\n')
    p.add_run('                let temp = arr[j];\n')
    p.add_run('                arr[j] = arr[j + 1];\n')
    p.add_run('                arr[j + 1] = temp;\n')
    p.add_run('            }\n')
    p.add_run('            j = j + 1;\n')
    p.add_run('        }\n')
    p.add_run('        i = i + 1;\n')
    p.add_run('    }\n')
    p.add_run('    return arr;\n')
    p.add_run('}\n')
    p.add_run('\n')
    p.add_run('# 示例 / Example\n')
    p.add_run('let numbers = [64, 34, 25, 12, 22, 11, 90];\n')
    p.add_run('bubble_sort(numbers);\n')
    p.add_run('print(numbers); # 输出: [11, 12, 22, 25, 34, 64, 90]')

    # A.8.2 快速排序 / Quick Sort
    p = doc.add_paragraph()
    run = p.add_run('A.8.2 快速排序 / Quick Sort')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('fn quick_sort(arr) {\n')
    p.add_run('    if (len(arr) < 2) { return arr; }\n')
    p.add_run('    let pivot = arr[0];\n')
    p.add_run('    let left = [];\n')
    p.add_run('    let right = [];\n')
    p.add_run('    let i = 1;\n')
    p.add_run('    while (i < len(arr)) {\n')
    p.add_run('        if (arr[i] < pivot) {\n')
    p.add_run('            push(left, arr[i]);\n')
    p.add_run('        } else {\n')
    p.add_run('            push(right, arr[i]);\n')
    p.add_run('        }\n')
    p.add_run('        i = i + 1;\n')
    p.add_run('    }\n')
    p.add_run('    return quick_sort(left) + [pivot] + quick_sort(right);\n')
    p.add_run('}\n')
    p.add_run('\n')
    p.add_run('# 示例 / Example\n')
    p.add_run('let arr = [33, 10, 55, 71, 29, 3];\n')
    p.add_run('let sorted = quick_sort(arr);\n')
    p.add_run('print(sorted); # 输出: [3, 10, 29, 33, 55, 71]')

    # A.9 查找算法 / Search Algorithms
    p = doc.add_paragraph()
    run = p.add_run('A.9 查找算法 / Search Algorithms')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(26, 86, 142)

    # A.9.1 线性查找 / Linear Search
    p = doc.add_paragraph()
    run = p.add_run('A.9.1 线性查找 / Linear Search')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('fn linear_search(arr, target) {\n')
    p.add_run('    let i = 0;\n')
    p.add_run('    while (i < len(arr)) {\n')
    p.add_run('        if (arr[i] == target) {\n')
    p.add_run('            return i;\n')
    p.add_run('        }\n')
    p.add_run('        i = i + 1;\n')
    p.add_run('    }\n')
    p.add_run('    return -1;\n')
    p.add_run('}\n')
    p.add_run('\n')
    p.add_run('# 示例 / Example\n')
    p.add_run('let data = [10, 20, 30, 40, 50];\n')
    p.add_run('let index = linear_search(data, 30);\n')
    p.add_run('print(index); # 输出: 2')

    # A.9.2 二分查找 / Binary Search
    p = doc.add_paragraph()
    run = p.add_run('A.9.2 二分查找 / Binary Search')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('fn binary_search(arr, target) {\n')
    p.add_run('    let left = 0;\n')
    p.add_run('    let right = len(arr) - 1;\n')
    p.add_run('    while (left <= right) {\n')
    p.add_run('        let mid = left + (right - left) / 2;\n')
    p.add_run('        if (arr[mid] == target) {\n')
    p.add_run('            return mid;\n')
    p.add_run('        }\n')
    p.add_run('        if (arr[mid] < target) {\n')
    p.add_run('            left = mid + 1;\n')
    p.add_run('        } else {\n')
    p.add_run('            right = mid - 1;\n')
    p.add_run('        }\n')
    p.add_run('    }\n')
    p.add_run('    return -1;\n')
    p.add_run('}\n')
    p.add_run('\n')
    p.add_run('# 示例 / Example (需要已排序数组)\n')
    p.add_run('let sorted = [1, 3, 5, 7, 9, 11, 13, 15];\n')
    p.add_run('let result = binary_search(sorted, 7);\n')
    p.add_run('print(result); # 输出: 3')

    # A.10 递归算法 / Recursive Algorithms
    p = doc.add_paragraph()
    run = p.add_run('A.10 递归算法 / Recursive Algorithms')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(26, 86, 142)

    # A.10.1 阶乘 / Factorial
    p = doc.add_paragraph()
    run = p.add_run('A.10.1 阶乘 / Factorial')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('fn factorial(n) {\n')
    p.add_run('    if (n <= 1) { return 1; }\n')
    p.add_run('    return n * factorial(n - 1);\n')
    p.add_run('}\n')
    p.add_run('\n')
    p.add_run('# 示例 / Example\n')
    p.add_run('print(factorial(5));  # 输出: 120\n')
    p.add_run('print(factorial(0));  # 输出: 1')

    # A.10.2 斐波那契数列 / Fibonacci
    p = doc.add_paragraph()
    run = p.add_run('A.10.2 斐波那契数列 / Fibonacci')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('# 递归版本 / Recursive Version\n')
    p.add_run('fn fibonacci(n) {\n')
    p.add_run('    if (n <= 1) { return n; }\n')
    p.add_run('    return fibonacci(n - 1) + fibonacci(n - 2);\n')
    p.add_run('}\n')
    p.add_run('\n')
    p.add_run('# 迭代版本（更高效）/ Iterative Version (more efficient)\n')
    p.add_run('fn fibonacci_iter(n) {\n')
    p.add_run('    if (n <= 1) { return n; }\n')
    p.add_run('    let a = 0;\n')
    p.add_run('    let b = 1;\n')
    p.add_run('    let i = 2;\n')
    p.add_run('    while (i <= n) {\n')
    p.add_run('        let temp = a + b;\n')
    p.add_run('        a = b;\n')
    p.add_run('        b = temp;\n')
    p.add_run('        i = i + 1;\n')
    p.add_run('    }\n')
    p.add_run('    return b;\n')
    p.add_run('}\n')
    p.add_run('\n')
    p.add_run('# 示例 / Example\n')
    p.add_run('let i = 0;\n')
    p.add_run('while (i < 10) {\n')
    p.add_run('    print(fibonacci_iter(i));\n')
    p.add_run('    i = i + 1;\n')
    p.add_run('}\n')
    p.add_run('# 输出: 0, 1, 1, 2, 3, 5, 8, 13, 21, 34')

    # A.11 字符串算法 / String Algorithms
    p = doc.add_paragraph()
    run = p.add_run('A.11 字符串算法 / String Algorithms')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(26, 86, 142)

    # A.11.1 回文检查 / Palindrome Check
    p = doc.add_paragraph()
    run = p.add_run('A.11.1 回文检查 / Palindrome Check')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('fn is_palindrome(s) {\n')
    p.add_run('    let left = 0;\n')
    p.add_run('    let right = len(s) - 1;\n')
    p.add_run('    while (left < right) {\n')
    p.add_run('        if (s[left] != s[right]) {\n')
    p.add_run('            return false;\n')
    p.add_run('        }\n')
    p.add_run('        left = left + 1;\n')
    p.add_run('        right = right - 1;\n')
    p.add_run('    }\n')
    p.add_run('    return true;\n')
    p.add_run('}\n')
    p.add_run('\n')
    p.add_run('# 示例 / Example\n')
    p.add_run('print(is_palindrome("radar"));   # 输出: true\n')
    p.add_run('print(is_palindrome("hello"));   # 输出: false')

    # A.11.2 字符串反转 / String Reverse
    p = doc.add_paragraph()
    run = p.add_run('A.11.2 字符串反转 / String Reverse')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('fn reverse_string(s) {\n')
    p.add_run('    let result = "";\n')
    p.add_run('    let i = len(s) - 1;\n')
    p.add_run('    while (i >= 0) {\n')
    p.add_run('        result = result + s[i];\n')
    p.add_run('        i = i - 1;\n')
    p.add_run('    }\n')
    p.add_run('    return result;\n')
    p.add_run('}\n')
    p.add_run('\n')
    p.add_run('# 示例 / Example\n')
    p.add_run('print(reverse_string("HSharp")); # 输出: praSH')

    # A.11.3 单词反转 / Word Reverse
    p = doc.add_paragraph()
    run = p.add_run('A.11.3 单词反转 / Word Reverse')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('fn reverse_words(s) {\n')
    p.add_run('    let words = [];\n')
    p.add_run('    let word = "";\n')
    p.add_run('    let i = 0;\n')
    p.add_run('    while (i < len(s)) {\n')
    p.add_run('        if (s[i] == " ") {\n')
    p.add_run('            if (len(word) > 0) {\n')
    p.add_run('                push(words, word);\n')
    p.add_run('                word = "";\n')
    p.add_run('            }\n')
    p.add_run('        } else {\n')
    p.add_run('            word = word + s[i];\n')
    p.add_run('        }\n')
    p.add_run('        i = i + 1;\n')
    p.add_run('    }\n')
    p.add_run('    if (len(word) > 0) {\n')
    p.add_run('        push(words, word);\n')
    p.add_run('    }\n')
    p.add_run('    \n')
    p.add_run('    let result = "";\n')
    p.add_run('    i = len(words) - 1;\n')
    p.add_run('    while (i >= 0) {\n')
    p.add_run('        result = result + words[i];\n')
    p.add_run('        if (i > 0) { result = result + " "; }\n')
    p.add_run('        i = i - 1;\n')
    p.add_run('    }\n')
    p.add_run('    return result;\n')
    p.add_run('}\n')
    p.add_run('\n')
    p.add_run('# 示例 / Example\n')
    p.add_run('print(reverse_words("Hello World")); # 输出: World Hello')

    # A.12 数学算法 / Mathematical Algorithms
    p = doc.add_paragraph()
    run = p.add_run('A.12 数学算法 / Mathematical Algorithms')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(26, 86, 142)

    # A.12.1 质数检查 / Prime Number Check
    p = doc.add_paragraph()
    run = p.add_run('A.12.1 质数检查 / Prime Number Check')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('fn is_prime(n) {\n')
    p.add_run('    if (n < 2) { return false; }\n')
    p.add_run('    if (n == 2) { return true; }\n')
    p.add_run('    if (n % 2 == 0) { return false; }\n')
    p.add_run('    let i = 3;\n')
    p.add_run('    while (i * i <= n) {\n')
    p.add_run('        if (n % i == 0) { return false; }\n')
    p.add_run('        i = i + 2;\n')
    p.add_run('    }\n')
    p.add_run('    return true;\n')
    p.add_run('}\n')
    p.add_run('\n')
    p.add_run('# 示例 / Example\n')
    p.add_run('print(is_prime(17));  # 输出: true\n')
    p.add_run('print(is_prime(18));  # 输出: false')

    # A.12.2 最大公约数 (GCD) / Greatest Common Divisor
    p = doc.add_paragraph()
    run = p.add_run('A.12.2 最大公约数 (GCD) / Greatest Common Divisor')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('# 欧几里得算法 / Euclidean Algorithm\n')
    p.add_run('fn gcd(a, b) {\n')
    p.add_run('    while (b != 0) {\n')
    p.add_run('        let temp = b;\n')
    p.add_run('        b = a % b;\n')
    p.add_run('        a = temp;\n')
    p.add_run('    }\n')
    p.add_run('    return a;\n')
    p.add_run('}\n')
    p.add_run('\n')
    p.add_run('# 示例 / Example\n')
    p.add_run('print(gcd(48, 18)); # 输出: 6\n')
    p.add_run('print(gcd(100, 25)); # 输出: 25')

    # A.12.3 最小公倍数 (LCM) / Least Common Multiple
    p = doc.add_paragraph()
    run = p.add_run('A.12.3 最小公倍数 (LCM) / Least Common Multiple')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('fn lcm(a, b) {\n')
    p.add_run('    return a * b / gcd(a, b);\n')
    p.add_run('}\n')
    p.add_run('\n')
    p.add_run('# 示例 / Example\n')
    p.add_run('print(lcm(12, 18)); # 输出: 36')

    # A.12.4 幂运算 / Power Operation
    p = doc.add_paragraph()
    run = p.add_run('A.12.4 幂运算 / Power Operation')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('# 迭代版本 / Iterative Version\n')
    p.add_run('fn power(base, exp) {\n')
    p.add_run('    if (exp == 0) { return 1; }\n')
    p.add_run('    let result = 1;\n')
    p.add_run('    let i = 0;\n')
    p.add_run('    while (i < exp) {\n')
    p.add_run('        result = result * base;\n')
    p.add_run('        i = i + 1;\n')
    p.add_run('    }\n')
    p.add_run('    return result;\n')
    p.add_run('}\n')
    p.add_run('\n')
    p.add_run('# 示例 / Example\n')
    p.add_run('print(power(2, 10));  # 输出: 1024\n')
    p.add_run('print(power(3, 4));   # 输出: 81')

    # A.13 综合示例 / Comprehensive Examples
    p = doc.add_paragraph()
    run = p.add_run('A.13 综合示例 / Comprehensive Examples')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(26, 86, 142)

    # A.13.1 数组扁平化 / Flatten Array
    p = doc.add_paragraph()
    run = p.add_run('A.13.1 数组扁平化 / Flatten Array')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('fn flatten(arr) {\n')
    p.add_run('    let result = [];\n')
    p.add_run('    let i = 0;\n')
    p.add_run('    while (i < len(arr)) {\n')
    p.add_run('        if (typeof(arr[i]) == "array") {\n')
    p.add_run('            let nested = flatten(arr[i]);\n')
    p.add_run('            let j = 0;\n')
    p.add_run('            while (j < len(nested)) {\n')
    p.add_run('                push(result, nested[j]);\n')
    p.add_run('                j = j + 1;\n')
    p.add_run('            }\n')
    p.add_run('        } else {\n')
    p.add_run('            push(result, arr[i]);\n')
    p.add_run('        }\n')
    p.add_run('        i = i + 1;\n')
    p.add_run('    }\n')
    p.add_run('    return result;\n')
    p.add_run('}\n')
    p.add_run('\n')
    p.add_run('# 示例 / Example\n')
    p.add_run('let nested = [1, [2, [3, 4]], 5, [6]];\n')
    p.add_run('print(flatten(nested)); # 输出: [1, 2, 3, 4, 5, 6]')

    # A.13.2 汉诺塔 / Tower of Hanoi
    p = doc.add_paragraph()
    run = p.add_run('A.13.2 汉诺塔 / Tower of Hanoi')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('fn hanoi(n, source, target, auxiliary) {\n')
    p.add_run('    if (n == 1) {\n')
    p.add_run('        print("Move disk 1 from " + source + " to " + target);\n')
    p.add_run('        return;\n')
    p.add_run('    }\n')
    p.add_run('    hanoi(n - 1, source, auxiliary, target);\n')
    p.add_run('    print("Move disk " + n + " from " + source + " to " + target);\n')
    p.add_run('    hanoi(n - 1, auxiliary, target, source);\n')
    p.add_run('}\n')
    p.add_run('\n')
    p.add_run('# 示例 / Example\n')
    p.add_run('hanoi(3, "A", "C", "B");\n')
    p.add_run('# 输出:\n')
    p.add_run('# Move disk 1 from A to C\n')
    p.add_run('# Move disk 2 from A to B\n')
    p.add_run('# Move disk 1 from C to B\n')
    p.add_run('# Move disk 3 from A to C\n')
    p.add_run('# Move disk 1 from B to A\n')
    p.add_run('# Move disk 2 from B to C\n')
    p.add_run('# Move disk 1 from A to C')

    # A.13.3 合并两个有序数组 / Merge Two Sorted Arrays
    p = doc.add_paragraph()
    run = p.add_run('A.13.3 合并两个有序数组 / Merge Two Sorted Arrays')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('fn merge_sorted(arr1, arr2) {\n')
    p.add_run('    let result = [];\n')
    p.add_run('    let i = 0;\n')
    p.add_run('    let j = 0;\n')
    p.add_run('    while (i < len(arr1) and j < len(arr2)) {\n')
    p.add_run('        if (arr1[i] < arr2[j]) {\n')
    p.add_run('            push(result, arr1[i]);\n')
    p.add_run('            i = i + 1;\n')
    p.add_run('        } else {\n')
    p.add_run('            push(result, arr2[j]);\n')
    p.add_run('            j = j + 1;\n')
    p.add_run('        }\n')
    p.add_run('    }\n')
    p.add_run('    while (i < len(arr1)) {\n')
    p.add_run('        push(result, arr1[i]);\n')
    p.add_run('        i = i + 1;\n')
    p.add_run('    }\n')
    p.add_run('    while (j < len(arr2)) {\n')
    p.add_run('        push(result, arr2[j]);\n')
    p.add_run('        j = j + 1;\n')
    p.add_run('    }\n')
    p.add_run('    return result;\n')
    p.add_run('}\n')
    p.add_run('\n')
    p.add_run('# 示例 / Example\n')
    p.add_run('let a = [1, 3, 5, 7];\n')
    p.add_run('let b = [2, 4, 6, 8];\n')
    p.add_run('print(merge_sorted(a, b)); # 输出: [1, 2, 3, 4, 5, 6, 7, 8]')

    # Save the document
    doc.save('/Users/peddlejumper/H#/v0.4/资料/HSharp_Programming_Tutorial.docx')
    print("Successfully added H# basic algorithm examples to Appendix!")

if __name__ == '__main__':
    add_hsharp_algorithm_examples()
