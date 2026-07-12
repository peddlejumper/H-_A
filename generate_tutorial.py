#!/usr/bin/env python3
"""Generate the H# Programming Tutorial in Word format (Chinese-English Bilingual)."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# ── Style helpers ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_heading_en(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        rPr.insert(0, rFonts)
    return h

def add_heading_cn(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        rPr.insert(0, rFonts)
    return h

def add_code_block(doc, code):
    p = doc.add_paragraph()
    try:
        p.style = doc.styles['CodeBlock']
    except KeyError:
        cb_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
        pf = cb_style.font
        pf.name = 'Consolas'
        pf.size = Pt(9.5)
        pf.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        p.style = cb_style
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.8)
    # background shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F0F0F5')
    shading.set(qn('w:val'), 'clear')
    pPr = p._element.get_or_add_pPr()
    pPr.append(shading)
    for line in code.split('\n'):
        if line == '':
            run = p.add_run('\n')
        else:
            run = p.add_run(line + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        rPr.insert(0, rFonts)
    return p

def add_para(doc, en_text, cn_text):
    """Add a paragraph with English then Chinese."""
    p = doc.add_paragraph()
    run_en = p.add_run(en_text)
    run_en.font.name = 'Calibri'
    run_en.font.size = Pt(11)
    rPr = run_en._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rPr.insert(0, rFonts)

    run_cn = p.add_run('\n' + cn_text)
    run_cn.font.name = 'Microsoft YaHei'
    run_cn.font.size = Pt(11)
    run_cn.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    rPr2 = run_cn._element.get_or_add_rPr()
    rFonts2 = OxmlElement('w:rFonts')
    rFonts2.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rPr2.insert(0, rFonts2)
    return p

def add_preface_para(doc, text, italic=False):
    """Add an English-only paragraph for the preface."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11.5)
    if italic:
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    else:
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rPr.insert(0, rFonts)
    return p

def add_para_en(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rPr.insert(0, rFonts)
    return p

def add_para_cn(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rPr.insert(0, rFonts)
    return p

def add_note(doc, en_text, cn_text):
    p = doc.add_paragraph()
    run_en = p.add_run('Note: ' + en_text)
    run_en.font.name = 'Calibri'
    run_en.font.size = Pt(10)
    run_en.font.italic = True
    run_en.font.color.rgb = RGBColor(0x88, 0x66, 0x00)
    rPr = run_en._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rPr.insert(0, rFonts)

    run_cn = p.add_run('\n注意：' + cn_text)
    run_cn.font.name = 'Microsoft YaHei'
    run_cn.font.size = Pt(10)
    run_cn.font.italic = True
    run_cn.font.color.rgb = RGBColor(0x88, 0x66, 0x00)
    rPr2 = run_cn._element.get_or_add_rPr()
    rFonts2 = OxmlElement('w:rFonts')
    rFonts2.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rPr2.insert(0, rFonts2)
    return p

# ═══════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_t = title.add_run('H# Programming Tutorial')
run_t.font.size = Pt(32)
run_t.font.bold = True
run_t.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)
run_t.font.name = 'Calibri'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_s = subtitle.add_run('H# 程序设计教程')
run_s.font.size = Pt(28)
run_s.font.bold = True
run_s.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
run_s.font.name = 'Microsoft YaHei'

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_v = ver.add_run('Version 0.4  |  Chinese-English Bilingual Edition  |  中英双语版')
run_v.font.size = Pt(14)
run_v.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run_v.font.name = 'Calibri'

doc.add_page_break()

# ═══════════════════════════════════════════════
#  TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════
add_heading_en(doc, 'Table of Contents / 目录', level=1)

toc_items = [
    ('Preface', 'Preface / 前言'),
    ('Chapter 1', 'Introduction / 简介'),
    ('Chapter 2', 'Getting Started / 快速入门'),
    ('Chapter 3', 'Basic Syntax / 基础语法'),
    ('Chapter 4', 'Variables and Data Types / 变量与数据类型'),
    ('Chapter 5', 'Operators / 运算符'),
    ('Chapter 6', 'Control Flow / 控制流'),
    ('Chapter 7', 'Functions / 函数'),
    ('Chapter 8', 'Arrays and Dictionaries / 数组与字典'),
    ('Chapter 9', 'Classes, Interfaces and Concepts / 类、接口与概念'),
    ('Chapter 10', 'Modules and Imports / 模块与导入'),
    ('Chapter 11', 'Error Handling / 错误处理'),
    ('Chapter 12', 'Asynchronous Programming / 异步编程'),
    ('Chapter 13', 'Standard Library Reference / 标准库参考'),
    ('Chapter 14', 'Self-Hosting Architecture / 自举架构'),
    ('Chapter 15', 'Spatial Region Programming / 空间区编程'),
    ('Chapter 16', 'zzwUI / HwdUI GUI Framework / zzwUI GUI框架'),
    ('Appendix A', 'Quick Reference / 快速参考'),
]

for num, name in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}    {name}')
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rPr.insert(0, rFonts)

doc.add_page_break()

# ═══════════════════════════════════════════════
#  PREFACE
# ═══════════════════════════════════════════════
add_heading_en(doc, 'Preface', level=1)

add_preface_para(doc,
    'There are moments in life when the ground beneath your feet simply vanishes. '
    'When everything you believed in — every promise, every whispered word, every '
    'shared dream — collapses into silence, leaving you standing alone in the wreckage '
    'of what you thought was real.'
)

add_preface_para(doc,
    'This book was born from one of those moments.'
)

add_preface_para(doc,
    'It was my third year at Zunyi No.12 Middle School — that delicate age when '
    'everything feels monumental and every glance carries the weight of the world. '
    'Like many teenagers, I had given my '
    'heart completely to someone, believing with the absolute certainty of youth '
    'that what we had was unbreakable. We walked the same school corridors, sat in the '
    'same classrooms, and built a world of private jokes and shared silences that '
    'I thought would last forever. I was wrong.'
)

add_preface_para(doc,
    'The end came without warning — or perhaps the warning signs were there all '
    'along, and I simply refused to see them. One ordinary afternoon, in an ordinary '
    'cafeteria, with ordinary words delivered in an extraordinarily cold tone, '
    'everything I had built my happiness upon was dismantled. No argument. No dramatic '
    'confrontation. Just a quiet, devastating finality that left me unable to breathe, '
    'unable to think, unable to feel anything except the crushing weight of absence.'
)

add_preface_para(doc,
    'The weeks that followed were a fog. I stopped paying attention in class. I stopped '
    'eating properly. I would lie in my bed at three in the morning, '
    'staring at the ceiling, replaying every conversation, every moment, trying to '
    'find the exact point where everything had gone wrong. The school that had once '
    'felt like a garden of possibility now felt like a prison of memories. Every '
    'corner held a ghost. Every song on my playlist was a wound. I was drowning in '
    'a sea of what-ifs, and I had no life raft.'
)

add_preface_para(doc,
    'But somewhere in that darkness, something unexpected began to happen.'
)

add_preface_para(doc,
    'One sleepless night, too exhausted to think about her anymore but too restless '
    'to sleep, I picked up a pen and a blank notebook. I did not plan to write anything '
    'in particular. I simply needed to do something — anything — to keep my hands '
    'busy and my mind from spiraling. I started sketching. Not drawings, but diagrams. '
    'Little boxes with arrows between them. A notation system for something I could '
    'not yet name. I wrote a few lines of what looked like code, a syntax that felt '
    'clean and honest, stripped of unnecessary ornamentation — perhaps because I myself '
    'felt stripped of everything unnecessary.'
)

add_preface_para(doc,
    'That first night, I filled twelve pages. The next night, twenty more. Within a '
    'week, the notebook was full. I bought another. And another. The act of creation '
    'became my survival mechanism. When the pain of memory threatened to overwhelm me, '
    'I would open my notebook and design. When loneliness pressed against my chest '
    'like a physical weight, I would write another specification, another module, '
    'another elegant rule for a language that was slowly taking shape under my pen.'
)

add_preface_para(doc,
    'By the end of that semester, I had written one hundred and fifty-five pages of '
    'design documents — entirely by hand, in black ink, on lined paper. A complete '
    'programming language specification: lexical grammar, syntax rules, type system, '
    'memory model, bytecode instruction set, standard library design, module system, '
    'error handling semantics, even a concurrency model. I had designed a self-hosting '
    'compiler architecture before I had ever written a single line of actual H# code. '
    'I named it H# — partly as a technical designation, partly as a quiet act of '
    'reclaiming something of my own.'
)

add_preface_para(doc,
    'Looking back now, I understand what I was really doing. In a world where '
    'everything had become unpredictable and painful, where human relationships '
    'had proven fragile and untrustworthy, I was building something that would '
    'never betray me. A system of absolute rules. A world where every statement '
    'had a defined meaning, where every function returned exactly what it promised, '
    'where cause and effect were transparent and reliable. I was constructing an '
    'orderly universe because the real one had collapsed around me.'
)

add_preface_para(doc,
    'This may sound melodramatic. It probably is. But it is also true — and I '
    'believe that many great works of engineering are born from similar places. '
    'We build not only to solve problems in the external world, but also to solve '
    'problems within ourselves. Every line of code can be a tiny act of healing. '
    'Every system design can be a way of making sense of chaos. The computer, in '
    'its cold and perfect logic, can become a sanctuary when the warm and imperfect '
    'world of human emotions has wounded us beyond what we think we can bear.'
)

add_preface_para(doc,
    'H# is not just a programming language. It is the artifact of a broken heart '
    'that refused to stay broken. It is one hundred and fifty-five pages of proof '
    'that creation is the most powerful response to destruction. It is my quiet '
    'answer to the question that haunted me during those dark months: "What do you '
    'do when everything you loved is gone?" You build something new. Something that '
    'is entirely, irrevocably yours.'
)

add_preface_para(doc,
    'Today, H# has grown far beyond those handwritten pages. It has a self-hosting '
    'compiler, a bytecode virtual machine, a complete standard library, a GUI '
    'framework, its own IDE, and a growing community. It has become something real, '
    'something useful, something that exists independently of the pain that created it. '
    'And that, perhaps, is the most important lesson of all: the things we build from '
    'our suffering do not have to remain defined by it. They can grow. They can evolve. '
    'They can become beautiful in their own right.'
)

add_preface_para(doc,
    'This tutorial is both a technical manual and a personal testament. If you are '
    'here to learn H#, I hope you find it clear, practical, and rigorous. If you are '
    'here because you are going through your own darkness, I hope you find something '
    'else: evidence that the most devastating chapters of our lives can become the '
    'prologue to something extraordinary. You are not alone. And the things you create '
    'in your hardest moments may turn out to be the most important things you ever build.'
)

# ── signature ──
p_sig = doc.add_paragraph()
p_sig.paragraph_format.space_before = Pt(24)
p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_sig = p_sig.add_run('— peddlejumper')
run_sig.font.name = 'Calibri'
run_sig.font.size = Pt(12)
run_sig.font.italic = True
run_sig.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
rPr = run_sig._element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
rPr.insert(0, rFonts)

p_place = doc.add_paragraph()
p_place.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_place = p_place.add_run('Zunyi No.12 Middle School, 2024')
run_place.font.name = 'Calibri'
run_place.font.size = Pt(10)
run_place.font.italic = True
run_place.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
rPr2 = run_place._element.get_or_add_rPr()
rFonts2 = OxmlElement('w:rFonts')
rFonts2.set(qn('w:eastAsia'), 'Microsoft YaHei')
rPr2.insert(0, rFonts2)

doc.add_page_break()

# ═══════════════════════════════════════════════
#  CHAPTER 1: INTRODUCTION
# ═══════════════════════════════════════════════
add_heading_en(doc, 'Chapter 1  Introduction', level=1)
add_heading_cn(doc, '第一章  简介', level=2)

add_para(doc,
    'H# is a modern, expressive programming language designed for simplicity, '
    'performance, and self-hosting capability. It features a clean C-like syntax '
    'with powerful functional programming concepts, making it suitable for '
    'systems programming, web development, and educational purposes. '
    'The language is dynamically typed and runs on a Python-based host runtime, '
    'with a self-hosting bytecode compiler and virtual machine written in H# itself.',
    'H# 是一门现代化、富有表现力的编程语言，设计目标为简洁、高性能和自举能力。'
    '它拥有类C的简洁语法，同时融合了强大的函数式编程概念，适用于系统编程、'
    'Web开发和教学用途。H# 是动态类型语言，运行在基于Python的主运行时之上，'
    '并拥有一个用H#自身编写的自举字节码编译器和虚拟机。'
)

add_para(doc,
    'Key features of H# include:',
    'H# 的主要特性包括：'
)

features = [
    ('Self-Hosting / 自举',
     'The H# compiler, bytecode executor, and interpreter are all written in H# itself.'),
    ('Dynamic Typing / 动态类型',
     'Variables do not require type declarations; types are resolved at runtime.'),
    ('First-Class Functions / 一等函数',
     'Functions can be assigned to variables, passed as arguments, and returned from other functions.'),
    ('Rich Standard Library / 丰富的标准库',
     'Built-in modules for math, strings, arrays, file I/O, networking, databases, and cryptography.'),
    ('Asynchronous Support / 异步支持',
     'Coroutine-based async programming with yield, sleep, and event signaling.'),
    ('Object-Oriented / 面向对象',
     'Classes, interfaces, inheritance, and the "concept" abstraction mechanism.'),
    ('Clean Syntax / 简洁语法',
     'C-style blocks, semicolons, and familiar operators with minimal boilerplate.'),
]

for en, cn in features:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(en)
    run.font.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rPr.insert(0, rFonts)
    run2 = p.add_run(' — ' + cn)
    run2.font.name = 'Microsoft YaHei'
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ═══════════════════════════════════════════════
#  CHAPTER 2: GETTING STARTED
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading_en(doc, 'Chapter 2  Getting Started', level=1)
add_heading_cn(doc, '第二章  快速入门', level=2)

add_heading_en(doc, '2.1  Your First H# Program', level=2)
add_heading_cn(doc, '2.1  第一个H#程序', level=3)

add_para(doc,
    'Create a file with the .hto extension (e.g., hello.hto) and write the following code:',
    '创建一个以 .hto 为扩展名的文件（例如 hello.hto），并写入以下代码：'
)

add_code_block(doc, 'print("Hello, world!");')

add_para(doc,
    'Run it using the H# interpreter:',
    '使用 H# 解释器运行：'
)

add_code_block(doc, 'python3 hsharp.py hello.hto')

add_para(doc,
    'Output:',
    '输出：'
)

add_code_block(doc, 'Hello, world!')

add_heading_en(doc, '2.2  Program Structure', level=2)
add_heading_cn(doc, '2.2  程序结构', level=3)

add_para(doc,
    'An H# program consists of a sequence of statements. Each statement typically '
    'ends with a semicolon (;). Statements can be variable declarations, '
    'expressions, function definitions, control flow blocks, imports, and more.',
    'H# 程序由一系列语句组成。每个语句通常以分号(;)结尾。'
    '语句可以是变量声明、表达式、函数定义、控制流块、导入等。'
)

add_code_block(doc, '''let x = 10;
let y = 20;
let sum = x + y;
print("The sum is " + sum);''')

add_note(doc,
    'Comments start with # and extend to the end of the line.',
    '注释以 # 开头，持续到行尾。'
)

# ═══════════════════════════════════════════════
#  CHAPTER 3: BASIC SYNTAX
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading_en(doc, 'Chapter 3  Basic Syntax', level=1)
add_heading_cn(doc, '第三章  基础语法', level=2)

add_heading_en(doc, '3.1  Statements and Semicolons', level=2)
add_heading_cn(doc, '3.1  语句和分号', level=3)

add_para(doc,
    'In H#, most statements end with a semicolon (;). The semicolon is required '
    'for statement termination. Block statements (inside {}) also require '
    'semicolons for each inner statement.',
    '在 H# 中，大多数语句以分号(;)结尾。分号是语句终止所必需的。'
    '块语句（{} 内部）内部的每个语句也需要分号。'
)

add_heading_en(doc, '3.2  Comments', level=2)
add_heading_cn(doc, '3.2  注释', level=3)

add_para(doc,
    'H# supports single-line comments using the # character. Everything from # '
    'to the end of the line is ignored by the interpreter.',
    'H# 支持使用 # 字符的单行注释。从 # 到行尾的所有内容都会被解释器忽略。'
)

add_code_block(doc, '''# This is a comment
let x = 42;  # This is an inline comment''')

add_heading_en(doc, '3.3  Identifiers', level=2)
add_heading_cn(doc, '3.3  标识符', level=3)

add_para(doc,
    'Identifiers consist of letters (a-z, A-Z), digits (0-9), and underscores (_). '
    'They must begin with a letter or underscore. Identifiers are case-sensitive.',
    '标识符由字母（a-z, A-Z）、数字（0-9）和下划线（_）组成。'
    '必须以字母或下划线开头。标识符区分大小写。'
)

add_heading_en(doc, '3.4  Keywords', level=2)
add_heading_cn(doc, '3.4  关键字', level=3)

add_para(doc,
    'The following words are reserved keywords in H# and cannot be used as identifiers:',
    '以下单词是 H# 中的保留关键字，不能用作标识符：'
)

add_code_block(doc, '''let    fn      return  while   if      else
for    in      print   import  class   new
extends private static  interface implements super
is      as      module  concept coro    asm
ptr     true    false   and     or      not
continue break   nullptr auto    try     catch
throw''')

# ═══════════════════════════════════════════════
#  CHAPTER 4: VARIABLES AND DATA TYPES
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading_en(doc, 'Chapter 4  Variables and Data Types', level=1)
add_heading_cn(doc, '第四章  变量与数据类型', level=2)

add_heading_en(doc, '4.1  Variable Declaration', level=2)
add_heading_cn(doc, '4.1  变量声明', level=3)

add_para(doc,
    'Variables in H# are declared using the let keyword. H# is dynamically typed, '
    'so you do not need to specify a type when declaring a variable.',
    'H# 中的变量使用 let 关键字声明。H# 是动态类型的，因此声明变量时无需指定类型。'
)

add_code_block(doc, '''let x = 42;
let name = "H#";
let flag = true;
let nothing = nullptr;''')

add_para(doc,
    'Once declared, variables can be reassigned new values (without let):',
    '声明后，可以为变量重新赋值（不需要 let）：'
)

add_code_block(doc, '''let x = 10;
x = 20;       # reassign
x = x + 5;    # x becomes 25''')

add_heading_en(doc, '4.2  Data Types', level=2)
add_heading_cn(doc, '4.2  数据类型', level=3)

add_para(doc,
    'H# supports the following built-in data types:',
    'H# 支持以下内置数据类型：'
)

types_info = [
    ('Number / 数字', 'Integer and floating-point numbers. Examples: 42, -7, 3.14'),
    ('String / 字符串', 'Text enclosed in double quotes. Examples: "hello", "H#"'),
    ('Boolean / 布尔值', 'Logical values: true and false'),
    ('Null / 空值', 'Represents the absence of a value: nullptr'),
    ('Array / 数组', 'Ordered collection of values: [1, 2, 3]'),
    ('Dictionary / 字典', 'Key-value pairs: [["name", "Alice"], ["age", 25]]'),
    ('Function / 函数', 'Callable function values (first-class)'),
]

for t_en, t_cn in types_info:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{t_en}')
    run.font.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run2 = p.add_run(f'\n{t_cn}')
    run2.font.name = 'Microsoft YaHei'
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

add_heading_en(doc, '4.3  String Operations', level=2)
add_heading_cn(doc, '4.3  字符串操作', level=3)

add_para(doc,
    'Strings can be concatenated with the + operator. Escape sequences like '
    '\\n (newline) and \\" (double quote) are supported.',
    '字符串可以使用 + 运算符进行拼接。支持转义序列，如 \\n（换行）和 \\"（双引号）。'
)

add_code_block(doc, '''let greeting = "Hello";
let name = "World";
let message = greeting + ", " + name + "!";
print(message);

let multi = "Line 1\\nLine 2";
print(multi);''')

add_heading_en(doc, '4.4  Truthiness', level=2)
add_heading_cn(doc, '4.4  真值判断', level=3)

add_para(doc,
    'In boolean contexts, false and nullptr are considered falsy. '
    'All other values (including 0, empty string "", empty array []) are considered truthy.',
    '在布尔上下文中，false 和 nullptr 被视为假值。'
    '所有其他值（包括 0、空字符串 ""、空数组 []）都被视为真值。'
)

# ═══════════════════════════════════════════════
#  CHAPTER 5: OPERATORS
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading_en(doc, 'Chapter 5  Operators', level=1)
add_heading_cn(doc, '第五章  运算符', level=2)

add_heading_en(doc, '5.1  Arithmetic Operators', level=2)
add_heading_cn(doc, '5.1  算术运算符', level=3)

add_para(doc,
    'H# supports standard arithmetic operators:',
    'H# 支持标准算术运算符：'
)

add_code_block(doc, '''let a = 10 + 3;   # Addition / 加法 → 13
let b = 10 - 3;   # Subtraction / 减法 → 7
let c = 10 * 3;   # Multiplication / 乘法 → 30
let d = 10 / 3;   # Division / 除法 → 3.333...
let e = -x;       # Unary negation / 一元取负''')

add_heading_en(doc, '5.2  Comparison Operators', level=2)
add_heading_cn(doc, '5.2  比较运算符', level=3)

add_code_block(doc, '''a == b    # Equal to / 等于
a != b    # Not equal to / 不等于
a > b     # Greater than / 大于
a < b     # Less than / 小于
a >= b    # Greater than or equal to / 大于等于
a <= b    # Less than or equal to / 小于等于''')

add_heading_en(doc, '5.3  Logical Operators', level=2)
add_heading_cn(doc, '5.3  逻辑运算符', level=3)

add_para(doc,
    'Logical operators use short-circuit evaluation:',
    '逻辑运算符使用短路求值：'
)

add_code_block(doc, '''true and false   # Logical AND / 逻辑与 → false
true or false    # Logical OR / 逻辑或 → true
not true         # Logical NOT / 逻辑非 → false''')

add_heading_en(doc, '5.4  Bitwise Operators', level=2)
add_heading_cn(doc, '5.4  位运算符', level=3)

add_code_block(doc, '''a & b    # Bitwise AND / 按位与
a | b    # Bitwise OR / 按位或
a ^ b    # Bitwise XOR / 按位异或
a << n   # Left shift / 左移
a >> n   # Right shift / 右移
~a       # Bitwise NOT / 按位取反''')

add_heading_en(doc, '5.5  Operator Precedence', level=2)
add_heading_cn(doc, '5.5  运算符优先级', level=3)

add_para(doc,
    'Operators follow standard mathematical precedence. Parentheses () '
    'can be used to explicitly control evaluation order:',
    '运算符遵循标准数学优先级。可以使用括号()显式控制求值顺序：'
)

add_code_block(doc, '''let x = 2 + 3 * 4;     # 14 (multiplication first)
let y = (2 + 3) * 4;   # 20 (parentheses first)''')

# ═══════════════════════════════════════════════
#  CHAPTER 6: CONTROL FLOW
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading_en(doc, 'Chapter 6  Control Flow', level=1)
add_heading_cn(doc, '第六章  控制流', level=2)

add_heading_en(doc, '6.1  If Statement', level=2)
add_heading_cn(doc, '6.1  If 语句', level=3)

add_code_block(doc, '''let x = 15;

if (x > 10) {
    print("x is greater than 10");
}

if (x > 10) {
    print("big");
} else {
    print("small");
}

if (x > 20) {
    print("large");
} else if (x > 10) {
    print("medium");
} else {
    print("small");
}''')

add_heading_en(doc, '6.2  While Loop', level=2)
add_heading_cn(doc, '6.2  While 循环', level=3)

add_code_block(doc, '''let i = 0;
while (i < 5) {
    print(i);
    i = i + 1;
}''')

add_heading_en(doc, '6.3  For-In Loop', level=2)
add_heading_cn(doc, '6.3  For-In 循环', level=3)

add_para(doc,
    'The for-in loop iterates over the elements of an array or string:',
    'For-in 循环用于遍历数组或字符串的元素：'
)

add_code_block(doc, '''let arr = [10, 20, 30, 40];
for (let item in arr) {
    print(item);
}

let name = "H#";
for (let ch in name) {
    print(ch);
}''')

add_para(doc,
    'For dictionaries and arrays of pairs, you can use two variables:',
    '对于字典和键值对数组，可以使用两个变量：'
)

add_code_block(doc, '''let pairs = [["a", 1], ["b", 2], ["c", 3]];
for (let key, val in pairs) {
    print(key + " = " + val);
}''')

add_heading_en(doc, '6.4  Break and Continue', level=2)
add_heading_cn(doc, '6.4  Break 和 Continue', level=3)

add_code_block(doc, '''let i = 0;
while (true) {
    i = i + 1;
    if (i > 5) { break; }
    if (i == 3) { continue; }
    print(i);
}
# prints: 1, 2, 4, 5''')

# ═══════════════════════════════════════════════
#  CHAPTER 7: FUNCTIONS
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading_en(doc, 'Chapter 7  Functions', level=1)
add_heading_cn(doc, '第七章  函数', level=2)

add_heading_en(doc, '7.1  Function Declaration', level=2)
add_heading_cn(doc, '7.1  函数声明', level=3)

add_para(doc,
    'Functions are declared using the fn keyword. The function body is enclosed in braces {}.',
    '函数使用 fn 关键字声明。函数体用花括号 {} 包裹。'
)

add_code_block(doc, '''fn greet(name) {
    print("Hello, " + name + "!");
}

greet("Alice");''')

add_heading_en(doc, '7.2  Return Values', level=2)
add_heading_cn(doc, '7.2  返回值', level=3)

add_code_block(doc, '''fn add(a, b) {
    return a + b;
}

let result = add(10, 20);
print(result);  # 30''')

add_heading_en(doc, '7.3  Multiple Parameters', level=2)
add_heading_cn(doc, '7.3  多参数', level=3)

add_code_block(doc, '''fn volume(width, height, depth) {
    return width * height * depth;
}

let v = volume(3, 4, 5);  # 60''')

add_heading_en(doc, '7.4  Recursion', level=2)
add_heading_cn(doc, '7.4  递归', level=3)

add_para(doc,
    'H# fully supports recursive function calls:',
    'H# 完全支持递归函数调用：'
)

add_code_block(doc, '''fn factorial(n) {
    if (n <= 1) { return 1; }
    return n * factorial(n - 1);
}

print(factorial(5));  # 120

fn fib(n) {
    if (n <= 1) { return n; }
    return fib(n - 1) + fib(n - 2);
}

print(fib(10));  # 55''')

add_heading_en(doc, '7.5  Nested Functions and Closures', level=2)
add_heading_cn(doc, '7.5  嵌套函数与闭包', level=3)

add_para(doc,
    'Functions can be defined inside other functions. Inner functions capture '
    'variables from their enclosing scope, forming closures.',
    '函数可以在其他函数内部定义。内部函数捕获外部作用域中的变量，形成闭包。'
)

add_code_block(doc, '''fn make_adder(x) {
    fn adder(y) {
        return x + y;
    }
    return adder;
}

let add5 = make_adder(5);
print(add5(10));  # 15''')

add_heading_en(doc, '7.6  Functions as Values', level=2)
add_heading_cn(doc, '7.6  函数作为值', level=3)

add_para(doc,
    'Functions are first-class values in H#. They can be assigned to variables, '
    'passed as arguments to other functions, and returned from functions.',
    '函数是 H# 中的一等值。它们可以赋值给变量、作为参数传递给其他函数、'
    '以及从函数中返回。'
)

add_code_block(doc, '''fn apply(func, x) {
    return func(x);
}

fn square(n) { return n * n; }

let result = apply(square, 5);  # 25
print(result);''')

# ═══════════════════════════════════════════════
#  CHAPTER 8: ARRAYS AND DICTIONARIES
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading_en(doc, 'Chapter 8  Arrays and Dictionaries', level=1)
add_heading_cn(doc, '第八章  数组与字典', level=2)

add_heading_en(doc, '8.1  Array Literals', level=2)
add_heading_cn(doc, '8.1  数组字面量', level=3)

add_code_block(doc, '''let empty = [];
let numbers = [1, 2, 3, 4, 5];
let mixed = [42, "hello", true, nullptr];''')

add_heading_en(doc, '8.2  Indexing', level=2)
add_heading_cn(doc, '8.2  索引访问', level=3)

add_code_block(doc, '''let arr = [10, 20, 30];
print(arr[0]);   # 10
print(arr[1]);   # 20

arr[2] = 99;
print(arr[2]);   # 99''')

add_heading_en(doc, '8.3  Dictionary Literals', level=2)
add_heading_cn(doc, '8.3  字典字面量', level=3)

add_para(doc,
    'Dictionaries in H# are represented as arrays of key-value pairs:',
    'H# 中的字典表示为键值对数组：'
)

add_code_block(doc, '''let person = [["name", "Alice"], ["age", 25], ["city", "Beijing"]];
let name = person[0];      # ["name", "Alice"]
print(name[1]);            # "Alice"
let age = person[1][1];    # 25''')

add_heading_en(doc, '8.4  Array Operations', level=2)
add_heading_cn(doc, '8.4  数组操作', level=3)

add_code_block(doc, '''let arr = [1, 2, 3];
let n = len(arr);      # 3
push(arr, 4);          # [1, 2, 3, 4]
let last = pop(arr);   # 4, arr is now [1, 2, 3]''')

add_heading_en(doc, '8.5  Member Access', level=2)
add_heading_cn(doc, '8.5  成员访问', level=3)

add_para(doc,
    'H# supports dot-notation member access for dictionary-like structures:',
    'H# 支持字典类结构的点号成员访问：'
)

add_code_block(doc, '''let obj = [["x", 10], ["y", 20]];
let val = obj.x;   # 10 (via MemberExpression)''')

add_note(doc,
    'Member access works on arrays of key-value pairs where the key matches the attribute name.',
    '成员访问适用于键值对数组，其中键与属性名匹配。'
)

# ═══════════════════════════════════════════════
#  CHAPTER 9: CLASSES AND CONCEPTS
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading_en(doc, 'Chapter 9  Classes, Interfaces and Concepts', level=1)
add_heading_cn(doc, '第九章  类、接口与概念', level=2)

add_heading_en(doc, '9.1  Class Declaration', level=2)
add_heading_cn(doc, '9.1  类声明', level=3)

add_para(doc,
    'H# supports object-oriented programming with classes. Classes can contain '
    'fields (with optional private modifier) and methods (functions defined inside the class body).',
    'H# 支持面向对象编程的类。类可以包含字段（可选的 private 修饰符）'
    '和方法（在类体内定义的函数）。'
)

add_code_block(doc, '''class Point {
    let x = 0;
    let y = 0;

    fn init(px, py) {
        x = px;
        y = py;
    }

    fn distance() {
        return x * x + y * y;
    }
}''')

add_heading_en(doc, '9.2  Creating Instances', level=2)
add_heading_cn(doc, '9.2  创建实例', level=3)

add_para(doc,
    'Use the new keyword to create instances of a class:',
    '使用 new 关键字创建类的实例：'
)

add_code_block(doc, '''let p = new Point(3, 4);
let d = p.distance();  # 25''')

add_heading_en(doc, '9.3  Inheritance', level=2)
add_heading_cn(doc, '9.3  继承', level=3)

add_para(doc,
    'Classes can extend a base class using the extends keyword:',
    '类可以使用 extends 关键字继承基类：'
)

add_code_block(doc, '''class Point3D extends Point {
    let z = 0;

    fn init(x, y, z) {
        super.init(x, y);
        this.z = z;
    }
}''')

add_heading_en(doc, '9.4  Interfaces', level=2)
add_heading_cn(doc, '9.4  接口', level=3)

add_code_block(doc, '''interface Drawable {
    fn draw();
    fn get_color();
}

class Circle implements Drawable {
    fn draw() {
        print("Drawing circle");
    }
    fn get_color() {
        return "red";
    }
}''')

add_heading_en(doc, '9.5  Static Members', level=2)
add_heading_cn(doc, '9.5  静态成员', level=3)

add_code_block(doc, '''class MathUtils {
    static fn pi() {
        return 3.14159;
    }
}

let p = MathUtils.pi();
print(p);''')

add_heading_en(doc, '9.6  Concepts', level=2)
add_heading_cn(doc, '9.6  概念（Concept）', level=3)

add_para(doc,
    'The concept keyword defines a higher-level abstraction that can include '
    'functions and other declarations. Concepts are H#\'s mechanism for '
    'organising code into reusable, composable units beyond traditional classes.',
    'concept 关键字定义了一个更高层次的抽象，可以包含函数和其他声明。'
    'Concept 是 H# 用于将代码组织成可重用、可组合单元（超越传统类）的机制。'
)

add_code_block(doc, '''concept Logger {
    fn log(msg) {
        print("[LOG] " + msg);
    }
    fn error(msg) {
        print("[ERROR] " + msg);
    }
}

Logger.log("Application started");''')

# ═══════════════════════════════════════════════
#  CHAPTER 10: MODULES AND IMPORTS
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading_en(doc, 'Chapter 10  Modules and Imports', level=1)
add_heading_cn(doc, '第十章  模块与导入', level=2)

add_heading_en(doc, '10.1  Importing Files', level=2)
add_heading_cn(doc, '10.1  导入文件', level=3)

add_para(doc,
    'Use the import statement to include code from other .hto files. '
    'The import path is relative to the current working directory.',
    '使用 import 语句引入其他 .hto 文件中的代码。'
    '导入路径相对于当前工作目录。'
)

add_code_block(doc, '''import "bootstrap/math_utils.hto";
import "bootstrap/string_utils.hto";

let result = math_utils.abs(-42);
print(result);''')

add_heading_en(doc, '10.2  Module Declaration', level=2)
add_heading_cn(doc, '10.2  模块声明', level=3)

add_para(doc,
    'You can create named modules using the module keyword:',
    '可以使用 module 关键字创建命名模块：'
)

add_code_block(doc, '''module mylib {
    fn greet() {
        return "Hello from mylib!";
    }
}

print(mylib.greet());''')

add_heading_en(doc, '10.3  Standard Library Modules', level=2)
add_heading_cn(doc, '10.3  标准库模块', level=3)

add_para(doc,
    'H# ships with a rich standard library organized into modules. '
    'See Chapter 13 for the full reference.',
    'H# 提供了丰富的标准库，按模块组织。完整参考请见第十三章。'
)

# ═══════════════════════════════════════════════
#  CHAPTER 11: ERROR HANDLING
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading_en(doc, 'Chapter 11  Error Handling', level=1)
add_heading_cn(doc, '第十一章  错误处理', level=2)

add_heading_en(doc, '11.1  Try-Catch', level=2)
add_heading_cn(doc, '11.1  Try-Catch', level=3)

add_para(doc,
    'H# supports try-catch for handling runtime errors:',
    'H# 支持 try-catch 来处理运行时错误：'
)

add_code_block(doc, '''try {
    let x = 10 / 0;
} catch (err) {
    print("Error: " + err);
}''')

add_heading_en(doc, '11.2  Throw', level=2)
add_heading_cn(doc, '11.2  Throw', level=3)

add_para(doc,
    'You can throw exceptions using the throw keyword:',
    '可以使用 throw 关键字抛出异常：'
)

add_code_block(doc, '''fn divide(a, b) {
    if (b == 0) {
        throw "Division by zero";
    }
    return a / b;
}

try {
    divide(10, 0);
} catch (e) {
    print("Caught: " + e);
}''')

# ═══════════════════════════════════════════════
#  CHAPTER 12: ASYNCHRONOUS PROGRAMMING
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading_en(doc, 'Chapter 12  Asynchronous Programming', level=1)
add_heading_cn(doc, '第十二章  异步编程', level=2)

add_heading_en(doc, '12.1  Coroutines', level=2)
add_heading_cn(doc, '12.1  协程', level=3)

add_para(doc,
    'H# provides coroutine-based asynchronous programming. A coroutine is '
    'declared with the coro keyword and can yield execution, sleep, and wait for events.',
    'H# 提供基于协程的异步编程。协程使用 coro 关键字声明，'
    '可以让出执行权、休眠和等待事件。'
)

add_code_block(doc, '''coro task1() {
    let i = 0;
    while (i < 3) {
        print("Task 1: " + i);
        coro_sleep(100);
        i = i + 1;
    }
}

coro task2() {
    let i = 0;
    while (i < 3) {
        print("Task 2: " + i);
        coro_sleep(150);
        i = i + 1;
    }
}

task1();
task2();
scheduler_run();''')

add_heading_en(doc, '12.2  Coroutine Functions', level=2)
add_heading_cn(doc, '12.2  协程函数', level=3)

add_para(doc,
    'Key coroutine functions:',
    '关键协程函数：'
)

add_code_block(doc, '''coro_yield()        # Yield execution / 让出执行权
coro_sleep(ms)       # Sleep for milliseconds / 休眠毫秒数
coro_wait(event)     # Wait for an event / 等待事件
coro_signal(event)   # Signal an event / 发送事件信号''')

# ═══════════════════════════════════════════════
#  CHAPTER 13: STANDARD LIBRARY REFERENCE
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading_en(doc, 'Chapter 13  Standard Library Reference', level=1)
add_heading_cn(doc, '第十三章  标准库参考', level=2)

add_para(doc,
    'This chapter provides a reference for H#\'s built-in standard library functions.',
    '本章提供 H# 内置标准库函数的参考。'
)

# ── Core Builtins ──
add_heading_en(doc, '13.1  Core Built-in Functions / 核心内置函数', level=2)

add_code_block(doc, '''len(x)              # Return length of array, string, or dict
                      # 返回数组、字符串或字典的长度
push(arr, val)      # Append value to array / 将值添加到数组末尾
pop(arr)            # Remove and return last element / 移除并返回最后一个元素
read_file(path)     # Read file contents as string / 读取文件内容为字符串
write_file(path, s) # Write string to file / 将字符串写入文件
input(prompt)       # Read line from stdin / 从标准输入读取一行
int(x)              # Convert to integer / 转换为整数
str(x)              # Convert to string / 转换为字符串
ord(ch)             # Get character code / 获取字符编码
chr(n)              # Get character from code / 从编码获取字符
substring(s, b, e)  # Extract substring / 提取子字符串
time_now()          # Current time in milliseconds / 当前时间（毫秒）''')

# ── Math ──
add_heading_en(doc, '13.2  Math Module / 数学模块', level=2)

add_code_block(doc, '''# Available via: import "bootstrap/math_utils.hto"
# or: import "bootstrap/math_extended.hto"

# Basic math operations
abs(x)              # Absolute value / 绝对值
min(a, b)           # Minimum / 最小值
max(a, b)           # Maximum / 最大值
pow(base, exp)      # Power / 幂运算
sqrt(x)             # Square root / 平方根
floor(x)            # Floor / 向下取整
ceil(x)             # Ceiling / 向上取整
round(x)            # Round to nearest integer / 四舍五入

# Trigonometric functions
sin(x), cos(x), tan(x)
asin(x), acos(x), atan(x)

# Constants
PI                  # 3.14159...
E                   # 2.71828...''')

# ── String ──
add_heading_en(doc, '13.3  String Module / 字符串模块', level=2)

add_code_block(doc, '''# Available via: import "bootstrap/string_utils.hto"

to_upper(s)         # Convert to uppercase / 转换为大写
to_lower(s)         # Convert to lowercase / 转换为小写
trim(s)             # Remove whitespace / 移除空白字符
starts_with(s, p)   # Check prefix / 检查前缀
ends_with(s, p)     # Check suffix / 检查后缀
contains(s, sub)    # Check substring / 检查子串
replace(s, old, n)  # Replace substring / 替换子串
split(s, delim)     # Split into array / 分割为数组
join(arr, delim)    # Join array to string / 将数组连接为字符串''')

# ── Array ──
add_heading_en(doc, '13.4  Array Module / 数组模块', level=2)

add_code_block(doc, '''# Available via: import "bootstrap/array_utils.hto"

len(arr)            # Array length / 数组长度
push(arr, v)        # Append / 追加
pop(arr)            # Remove last / 弹出最后一个
index_of(arr, v)    # Find index / 查找索引
contains(arr, v)    # Check existence / 检查是否存在
slice(arr, b, e)    # Sub-array / 子数组
reverse(arr)        # Reverse in place / 原地反转
sort(arr)           # Sort in place / 原地排序
map(arr, fn)        # Apply function / 映射函数
filter(arr, fn)     # Filter elements / 过滤元素
reduce(arr, fn, i)  # Reduce / 归约''')

# ── File System ──
add_heading_en(doc, '13.5  File System Module / 文件系统模块', level=2)

add_code_block(doc, '''# Available via: import "bootstrap/fs_module.hto"

fs_exists(path)         # Check if path exists / 检查路径是否存在
fs_is_file(path)        # Check if path is a file / 检查是否为文件
fs_is_dir(path)         # Check if path is a directory / 检查是否为目录
fs_mkdir(path)          # Create directory / 创建目录
fs_remove(path)         # Remove file or directory / 删除文件或目录
fs_list_dir(path)       # List directory contents / 列出目录内容
fs_get_cwd()            # Get current working directory / 获取当前工作目录
fs_chdir(path)          # Change working directory / 更改工作目录
fs_join_path(a, b)      # Join path components / 连接路径组件
fs_get_ext(path)        # Get file extension / 获取文件扩展名
fs_get_basename(path)   # Get base filename / 获取基本文件名
fs_get_dirname(path)    # Get directory name / 获取目录名''')

# ── I/O ──
add_heading_en(doc, '13.6  I/O Module / IO 模块', level=2)

add_code_block(doc, '''# Available via: import "bootstrap/io_module.hto"

write_file(path, content)   # Write file / 写文件
read_file(path)             # Read file / 读文件
io_append_file(path, data)  # Append to file / 追加到文件
io_read_lines(path)         # Read file as lines / 按行读取文件
io_write_lines(path, lines) # Write lines to file / 将行数组写入文件''')

# ── Network ──
add_heading_en(doc, '13.7  Network Module / 网络模块', level=2)

add_code_block(doc, '''# Available via: import "bootstrap/net_module.hto"

http_get(url)               # HTTP GET request / HTTP GET 请求
http_post(url, data)        # HTTP POST request / HTTP POST 请求
url_parse(url)               # Parse URL / 解析 URL
url_build(parts)             # Build URL / 构建 URL
tcp_connect(host, port)      # TCP connect / TCP 连接
tcp_send(sock, data)         # TCP send / TCP 发送
tcp_recv(sock, size)         # TCP receive / TCP 接收
tcp_close(sock)              # TCP close / TCP 关闭
udp_create()                 # UDP socket / UDP 套接字
udp_send(sock, host, port, d)# UDP send / UDP 发送
udp_recv(sock, size)         # UDP receive / UDP 接收
base64_encode(data)          # Base64 encode / Base64 编码
base64_decode(data)          # Base64 decode / Base64 解码
json_stringify(obj)          # JSON encode / JSON 编码
json_parse(s)                # JSON decode / JSON 解码''')

# ── Database ──
add_heading_en(doc, '13.8  Database Module / 数据库模块', level=2)

add_code_block(doc, '''# Available via: import "bootstrap/db_module.hto"

db_connect(path)                # Connect to database / 连接数据库
db_close(conn)                  # Close connection / 关闭连接
db_execute(conn, sql)          # Execute SQL / 执行 SQL
db_query(conn, sql)            # Query rows / 查询结果行
db_query_one(conn, sql)        # Query single row / 查询单行
db_begin_transaction(conn)     # Begin transaction / 开始事务
db_commit(conn)                # Commit / 提交
db_rollback(conn)              # Rollback / 回滚
db_create_table(conn, name, c) # Create table / 创建表
db_drop_table(conn, name)      # Drop table / 删除表
db_get_tables(conn)            # List tables / 列出表
db_get_table_info(conn, name)  # Table info / 表信息''')

# ── Date/Time ──
add_heading_en(doc, '13.9  Date/Time Module / 日期时间模块', level=2)

add_code_block(doc, '''# Available via: import "bootstrap/datetime_module.hto"

date_now()                  # Current date / 当前日期
date_timestamp()            # Current timestamp / 当前时间戳
date_format(ts, fmt)        # Format timestamp / 格式化时间戳
date_parse(s, fmt)          # Parse date string / 解析日期字符串''')

# ── Crypto ──
add_heading_en(doc, '13.10  Cryptography Module / 加密模块', level=2)

add_code_block(doc, '''# Available via: import "bootstrap/crypto_module.hto"

# Hash functions
sha256(data)            # SHA-256 hash
md5(data)               # MD5 hash

# Encryption / decryption
aes_encrypt(data, key)  # AES encryption / AES 加密
aes_decrypt(data, key)  # AES decryption / AES 解密

# Encoding
hex_encode(data)        # Hex encode / 十六进制编码
hex_decode(data)        # Hex decode / 十六进制解码''')

# ═══════════════════════════════════════════════
#  CHAPTER 14: SELF-HOSTING
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading_en(doc, 'Chapter 14  Self-Hosting Architecture', level=1)
add_heading_cn(doc, '第十四章  自举架构', level=2)

add_heading_en(doc, '14.1  The Self-Hosting Pipeline', level=2)
add_heading_cn(doc, '14.1  自举流水线', level=3)

add_para(doc,
    'One of the most remarkable features of H# is its self-hosting capability. '
    'The core language toolchain — tokenizer, parser, compiler, and bytecode executor — '
    'are all written in H# itself. These components run on a minimal Python host '
    'that provides only basic I/O and data structure operations.',
    'H# 最显著的特性之一是其自举能力。核心语言工具链——词法分析器、解析器、'
    '编译器和字节码执行器——全部用 H# 自身编写。这些组件运行在一个极简的 Python '
    '宿主之上，该宿主仅提供基本的 I/O 和数据结构操作。'
)

add_para(doc,
    'The self-hosting pipeline consists of four stages:',
    '自举流水线包含四个阶段：'
)

add_code_block(doc, '''╔═══════════╗    ╔═════════╗    ╔══════════╗    ╔══════════╗
║ Tokenizer ║ →  ║ Parser  ║ →  ║ Compiler ║ →  ║ Executor ║
║ tokenize  ║    ║ parse   ║    ║ compile  ║    ║ run      ║
║  .hto     ║    ║  .hto   ║    ║  .hto    ║    ║  .hto    ║
╚═══════════╝    ╚═════════╝    ╚══════════╝    ╚══════════╝
     ↓                ↓               ↓               ↓
  Tokens           AST           Bytecode         Result''')

add_para(doc,
    'There is also a tree-walking interpreter (interpreter.hto) that evaluates '
    'AST nodes directly without compilation, providing an alternative execution path '
    'useful for development, debugging, and comparison testing.',
    '此外还有一个树遍历解释器（interpreter.hto），它直接求值 AST 节点而无需编译，'
    '为开发、调试和对比测试提供了另一条执行路径。'
)

add_heading_en(doc, '14.2  Tokenizer (tokenize.hto)', level=2)
add_heading_cn(doc, '14.2  词法分析器 (tokenize.hto)', level=3)

add_para(doc,
    'The H# tokenizer converts raw source code into a list of tokens. Each token '
    'is represented as a pair [type, value]. It handles keywords, identifiers, '
    'numbers, strings, operators, and punctuation.',
    'H# 词法分析器将原始源代码转换为 token 列表。每个 token 表示为一个 '
    '[类型, 值] 对。它处理关键字、标识符、数字、字符串、运算符和标点符号。'
)

add_heading_en(doc, '14.3  Parser (parser.hto)', level=2)
add_heading_cn(doc, '14.3  解析器 (parser.hto)', level=3)

add_para(doc,
    'The H# parser implements a recursive descent parser that converts the token '
    'stream into an Abstract Syntax Tree (AST). The AST is represented as nested '
    'arrays using a Lisp-like S-expression format. For example, "1 + 2" becomes '
    '["BinaryOp", ["NumberLiteral", "1"], "PLUS", ["NumberLiteral", "2"]].',
    'H# 解析器实现了一个递归下降解析器，将 token 流转换为抽象语法树（AST）。'
    'AST 使用类似 Lisp 的 S-表达式格式表示为嵌套数组。'
    '例如，"1 + 2" 变为 ["BinaryOp", ["NumberLiteral", "1"], "PLUS", ["NumberLiteral", "2"]]。'
)

add_heading_en(doc, '14.4  Compiler (compiler.hto)', level=2)
add_heading_cn(doc, '14.4  编译器 (compiler.hto)', level=3)

add_para(doc,
    'The H# compiler walks the AST and generates bytecode instructions for '
    'a stack-based virtual machine. It maintains a constants pool and emits '
    'instructions like LOAD_CONST, STORE_NAME, BINARY_ADD, CALL_FUNCTION, '
    'JUMP_IF_FALSE, etc. Functions are compiled into separate bytecode objects '
    'stored in the constants pool.',
    'H# 编译器遍历 AST，为基于栈的虚拟机生成字节码指令。'
    '它维护一个常量池，并发出 LOAD_CONST、STORE_NAME、BINARY_ADD、'
    'CALL_FUNCTION、JUMP_IF_FALSE 等指令。函数被编译为存储在常量池中的独立字节码对象。'
)

add_heading_en(doc, '14.5  Executor (executor.hto)', level=2)
add_heading_cn(doc, '14.5  执行器 (executor.hto)', level=3)

add_para(doc,
    'The bytecode executor is a stack-based VM that executes the compiled bytecode. '
    'It supports operand stack operations, environment (variable scope) management, '
    'function calls with closure support, and recursive calls.',
    '字节码执行器是一个基于栈的虚拟机，执行编译后的字节码。'
    '它支持操作数栈操作、环境（变量作用域）管理、带闭包支持的函数调用和递归调用。'
)

add_heading_en(doc, '14.6  Interpreter (interpreter.hto)', level=2)
add_heading_cn(doc, '14.6  解释器 (interpreter.hto)', level=3)

add_para(doc,
    'The tree-walking interpreter provides an alternative execution mode. It '
    'evaluates AST nodes directly without generating bytecode. This is useful '
    'for interactive development, testing, and as a reference implementation '
    'to verify the correctness of the compiler and executor.',
    '树遍历解释器提供了另一种执行模式。它直接求值 AST 节点而无需生成字节码。'
    '这对于交互式开发、测试以及作为验证编译器和执行器正确性的参考实现非常有用。'
)

add_heading_en(doc, '14.7  Verification', level=2)
add_heading_cn(doc, '14.7  验证', level=3)

add_para(doc,
    'The entire self-hosting chain is verified through comprehensive test suites '
    '(demo_executor.hto, test_bootstrap_full.hto) that execute code through both '
    'the bytecode executor and the tree-walking interpreter, ensuring consistent results.',
    '整个自举链通过全面的测试套件（demo_executor.hto, test_bootstrap_full.hto）进行验证，'
    '测试同时使用字节码执行器和树遍历解释器执行代码，确保结果一致。'
)

# ═══════════════════════════════════════════════
#  CHAPTER 15: SPATIAL REGION PROGRAMMING
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading_en(doc, 'Chapter 15  Spatial Region Programming', level=1)
add_heading_cn(doc, '第十五章  空间区编程', level=2)

add_heading_en(doc, '15.1  Overview', level=2)
add_heading_cn(doc, '15.1  概述', level=3)

add_para(doc,
    'H# introduces a unique paradigm called Spatial Region Programming (空间区编程), '
    'which allows developers to model and program three-dimensional spaces as '
    'first-class language constructs. This is achieved through a set of dedicated '
    'keywords: 3dsizepower, em3d, region, and region_interface. '
    'This paradigm is particularly suited for 3D game engines, spatial simulations, '
    'CAD systems, VR/AR applications, and any domain requiring explicit 3D spatial reasoning.',
    'H# 引入了一种独特的编程范式——空间区编程，允许开发者将三维空间作为一等语言结构来建模和编程。'
    '这通过一组专用关键字实现：3dsizepower、em3d、region 和 region_interface。'
    '此范式特别适用于3D游戏引擎、空间模拟、CAD系统、VR/AR应用，'
    '以及任何需要显式三维空间推理的领域。'
)

add_para(doc,
    'The spatial region system is built on four core concepts:',
    '空间区系统建立在四个核心概念之上：'
)

space_concepts = [
    ('3dsizepower / 三维大小力度空间',
     'A top-level 3D coordinate space container. It defines the spatial coordinate system, '
     'hold properties like dimensions and scales, and contains regions and region interfaces.\n'
     '顶层3D坐标空间容器。它定义空间坐标系，持有维度、尺度等属性，并包含区域和区域接口。'),
    ('region / 区域',
     'A bounded 3D sub-space defined by coordinate ranges (x1, y1, z1, x2, y2, z2). '
     'Each region can contain functions, classes, and methods that respond to events within its boundaries.\n'
     '由坐标范围 (x1, y1, z1, x2, y2, z2) 界定的3D子空间。'
     '每个区域可包含在其边界内响应事件的函数、类和方法。'),
    ('region_interface / 区域接口',
     'Defines a behavioral contract for regions, specifying methods that regions must implement. '
     'Like a traditional interface but tailored for spatial entities.\n'
     '定义区域的行为约定，指定区域必须实现的方法。类似传统接口，但专为空间实体定制。'),
    ('em3d / 扩展三维空间',
     'Extends a 3dsizepower, inheriting its regions, interfaces, and properties. '
     'Acts as a specialization or extension of an existing spatial coordinate system.\n'
     '继承自3dsizepower，继承其区域、接口和属性。'
     '作为现有空间坐标系统的特化或扩展。'),
]

for en, cn in space_concepts:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(en)
    run.font.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rPr.insert(0, rFonts)
    run2 = p.add_run(' — ' + cn)
    run2.font.name = 'Microsoft YaHei'
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

add_heading_en(doc, '15.2  Defining a 3D Coordinate Space', level=2)
add_heading_cn(doc, '15.2  定义三维坐标空间', level=3)

add_para(doc,
    'The 3dsizepower keyword defines a spatial coordinate system. Inside it, you can define '
    'public properties (using the public keyword), measurement attributes, regions, and region interfaces. '
    'Properties can use curly-brace {} syntax for parameter lists, square-bracket [] syntax for '
    'coordinate expressions (including auto{} for automatic parameter derivation), '
    'or equal-sign = syntax for expression-assigned values.',
    '3dsizepower 关键字定义了一个空间坐标系统。在其内部，可以定义 '
    '公有属性（使用 public 关键字）、度量属性、区域和区域接口。'
    '属性可使用花括号 {} 语法表示参数列表，方括号 [] 语法表示坐标表达式'
    '（包括 auto{} 自动参数推导），或等号 = 语法表示表达式赋值。'
)

add_code_block(doc, '''3dsizepower GameWorld {
    # Public properties with parameter lists
    public abs{T, T, size}
    public rel{playerSpawn, delta}

    # Coordinate expression with auto-derivation
    public size[auto{size=auto^3d}]

    # Scale factor as an expression
    public scale_factor = 1.0

    # Region definitions
    region SpawnZone(0, 0, 0, 100, 100, 100) {
        fn on_enter(entity) {
            print("Entity entered spawn zone");
        }
    }

    region BattleZone(100, 0, 0, 500, 300, 200) {
        class Enemy {
            fn attack() {
                print("Enemy attacks!");
            }
        }
    }
}''')

add_heading_en(doc, '15.3  Region Declaration', level=2)
add_heading_cn(doc, '15.3  区域声明', level=3)

add_para(doc,
    'A region is declared with the region keyword followed by a name and coordinate parameters '
    '(x1, y1, z1, x2, y2, z2). The six numbers define the bounding box of the region in 3D space. '
    'Inside the region body, you can define functions, classes, and methods. '
    'Regions can also implement region interfaces using the implements keyword.',
    '使用 region 关键字声明区域，后跟名称和坐标参数 (x1, y1, z1, x2, y2, z2)。'
    '六个数字定义了该区域在3D空间中的边界框。'
    '在区域体内，可以定义函数、类和方法。'
    '区域还可以使用 implements 关键字实现区域接口。'
)

add_code_block(doc, '''# Simple region
region Forest(0, 0, 0, 256, 64, 256) {
    fn ambient() {
        print("You hear birds chirping");
    }
    fn spawn_tree(count) {
        let i = 0;
        while (i < count) {
            print("Tree " + i + " spawned");
            i = i + 1;
        }
    }
}

# Region with interface implementation
region_interface Interactive {
    fn on_click(x, y, z);
    fn on_hover(x, y, z);
}

region TreasureRoom(50, 0, 50, 150, 30, 150)
    implements Interactive {
    fn on_click(x, y, z) {
        print("Treasure chest opened at " +
              x + "," + y + "," + z);
    }
    fn on_hover(x, y, z) {
        print("You see a glint of gold");
    }
}''')

add_heading_en(doc, '15.4  Region Interface', level=2)
add_heading_cn(doc, '15.4  区域接口', level=3)

add_para(doc,
    'A region_interface defines method signatures that regions can implement. It supports '
    'inheritance via the extends keyword, allowing one interface to extend another. '
    'This enables polymorphic behavior across different spatial regions.',
    'region_interface 定义区域可以实现的方法签名。它支持通过 extends 关键字进行继承，'
    '允许一个接口扩展另一个接口。这使得不同空间区域之间可以实现多态行为。'
)

add_code_block(doc, '''region_interface Damageable {
    fn take_damage(amount);
    fn get_health();
}

region_interface InteractiveDamageable
    extends Damageable {
    fn on_interact(player);
    fn highlight();
}

region BossArena(200, 0, 200, 500, 100, 500)
    implements InteractiveDamageable {
    fn take_damage(amount) {
        print("Boss took " + amount + " damage");
    }
    fn get_health() { return 1000; }
    fn on_interact(player) {
        print(player + " challenges the boss!");
    }
    fn highlight() {
        print("Boss arena glows red");
    }
}''')

add_heading_en(doc, '15.5  em3d — Extended 3D Space', level=2)
add_heading_cn(doc, '15.5  em3d — 扩展三维空间', level=3)

add_para(doc,
    'The em3d keyword extends an existing 3dsizepower, inheriting all its regions, '
    'interfaces, and properties. This is useful for creating specialized versions of '
    'a spatial system — for example, a night version of a game world, a higher-resolution '
    'simulation grid, or a subspace of a larger coordinate system.',
    'em3d 关键字扩展了一个已有的 3dsizepower，继承其所有区域、接口和属性。'
    '这对于创建空间系统的特化版本非常有用——例如，游戏世界的夜间版本、'
    '高分辨率模拟网格、或大型坐标系统的子空间。'
)

add_code_block(doc, '''3dsizepower BaseWorld {
    public abs{T, T, size}
    public size[auto{size=auto^3d}]

    region DefaultZone(0, 0, 0, 1000, 500, 1000) {
        fn basic_ambient() {
            print("Base world ambient");
        }
    }
}

em3d NightWorld extends BaseWorld {
    public lighting_level = 0.2
    public ambient_color = "dark blue"

    region ShadowZone(200, 0, 200, 400, 100, 400) {
        fn special_ambient() {
            print("Eerie shadows loom...");
        }
    }
}''')

add_heading_en(doc, '15.6  Working with 3D Points and Vectors', level=2)
add_heading_cn(doc, '15.6  使用3D点和向量', level=3)

add_para(doc,
    'H# represents 3D points and vectors as arrays of three numbers [x, y, z]. '
    'The d3system.hto module provides a comprehensive set of functions for 3D math operations:',
    'H# 将3D点和向量表示为三个数字的数组 [x, y, z]。'
    'd3system.hto 模块提供了一套全面的3D数学运算函数：'
)

add_code_block(doc, '''D3Point(x, y, z)          # Create a point / 创建点
D3Vec3(x, y, z)          # Create a vector / 创建向量

# Point operations / 点运算
d3_point_x(point)        # Get x coordinate / 获取x坐标
d3_point_y(point)        # Get y coordinate / 获取y坐标
d3_point_z(point)        # Get z coordinate / 获取z坐标
d3_point_eq(a, b)        # Point equality / 点相等判断
d3_point_add(a, b)       # Point addition / 点加法
d3_point_sub(a, b)       # Point subtraction / 点减法
d3_point_dist(a, b)      # Euclidean distance / 欧几里得距离
d3_point_mid(a, b)       # Midpoint / 中点

# Vector operations / 向量运算
d3_vec3_dot(a, b)        # Dot product / 点积
d3_vec3_cross(a, b)      # Cross product / 叉积
d3_vec3_length(v)        # Vector length / 向量长度
d3_vec3_normalize(v)     # Normalize vector / 归一化向量''')

add_heading_en(doc, '15.7  Bounding Box Operations', level=2)
add_heading_cn(doc, '15.7  边界框操作', level=3)

add_para(doc,
    'Bounding boxes are represented as six-element arrays [minX, minY, minZ, maxX, maxY, maxZ]. '
    'The D3 system provides functions for spatial queries:',
    '边界框表示为六元素数组 [minX, minY, minZ, maxX, maxY, maxZ]。'
    'D3系统提供了空间查询函数：'
)

add_code_block(doc, '''D3BoundingBox(x1,y1,z1,x2,y2,z2)  # Create bbox / 创建边界框
d3_bbox_min(box)          # Get minimum corner / 获取最小角点
d3_bbox_max(box)          # Get maximum corner / 获取最大角点
d3_bbox_center(box)       # Get center point / 获取中心点
d3_bbox_volume(box)       # Calculate volume / 计算体积
d3_bbox_contains(box, p)  # Point inside? / 点是否在内部？
d3_bbox_intersects(a, b)  # Boxes overlap? / 边界框是否重叠？
d3_bbox_intersection(a,b) # Compute overlap / 计算重叠区域
d3_point_in_range(p, x1,y1,z1,x2,y2,z2)  # Point in range?''')

add_heading_en(doc, '15.8  Region Lookup and Spatial Queries', level=2)
add_heading_cn(doc, '15.8  区域查找与空间查询', level=3)

add_para(doc,
    'The D3 system provides functions to query regions within a 3D coordinate space. '
    'These functions enable spatial reasoning such as finding which region contains a point, '
    'retrieving regions by name, and finding regions that intersect a bounding volume.',
    'D3系统提供了在3D坐标空间内查询区域的函数。'
    '这些函数支持空间推理，例如查找包含某个点的区域、按名称检索区域、'
    '以及查找与某个体积相交的区域。'
)

add_code_block(doc, '''# Find region containing a point / 查找包含某点的区域
let zone = d3_find_region_by_point(GameWorld, 50, 50, 50);
if (len(zone) > 0) {
    print("You are in: " + zone.name);
}

# Find region by name / 按名称查找区域
let battle = d3_find_region_by_name(GameWorld, "BattleZone");

# Get regions intersecting a bounding volume / 获取与边界体积相交的区域
let hits = d3_get_regions_intersecting(
    GameWorld, 0, 0, 0, 200, 150, 200);
let i = 0;
while (i < len(hits)) {
    print("Intersecting: " + hits[i].name);
    i = i + 1;
}''')

add_heading_en(doc, '15.9  3D Transforms', level=2)
add_heading_cn(doc, '15.9  3D变换', level=3)

add_para(doc,
    'H# supports basic 3D spatial transformations including translation, rotation, and scaling:',
    'H# 支持基本的3D空间变换，包括平移、旋转和缩放：'
)

add_code_block(doc, '''# Translate(3) + Rotation(axis,angle) + Scale(3)
D3Transform(tx, ty, tz, rotAxis, rotAngle, sx, sy, sz)

d3_transform_point(point, transform)
# Applies translation + scaling to a point''')

add_heading_en(doc, '15.10  System Utilities', level=2)
add_heading_cn(doc, '15.10  系统工具', level=3)

add_para(doc,
    'Convenience functions for inspecting and exporting spatial data:',
    '用于检查和导出空间数据的便捷函数：'
)

add_code_block(doc, '''d3_size_power_get_info(d3Obj)
# Prints summary: region count, interface count
# 打印摘要：区域数、接口数

d3_system_summary(d3Obj)
# Prints region count and total volume, returns volume
# 打印区域数和总体积，返回体积

d3_export_region_data(d3Obj, regionName)
# Returns [x1,y1,z1,x2,y2,z2] for a named region
# 返回命名区域的坐标数据 [x1,y1,z1,x2,y2,z2]''')

add_heading_en(doc, '15.11  Complete Example', level=2)
add_heading_cn(doc, '15.11  完整示例', level=3)

add_para(doc,
    'A complete spatial region program that defines a world, queries regions, '
    'and performs spatial calculations:',
    '一个完整的空间区程序，定义了一个世界、查询区域并执行空间计算：'
)

add_code_block(doc, '''import "bootstrap/d3system.hto";

3dsizepower MedievalWorld {
    public abs{T, T, size}
    public size[auto{size=auto^3d}]
    public scale_factor = 1.0

    region Castle(0, 0, 0, 200, 150, 200) {
        fn describe() {
            print("A majestic stone castle stands before you");
        }
        class Guard {
            fn patrol() {
                print("Guard patrols the ramparts");
            }
        }
    }

    region Village(200, 0, 0, 500, 30, 300) {
        fn describe() {
            print("A peaceful village with thatched cottages");
        }
    }

    region Forest(0, 0, 200, 500, 80, 500) {
        fn describe() {
            print("A dense, dark forest");
        }
    }
}

# Query the world
d3_size_power_get_info(MedievalWorld);

# Find player location
let playerPos = [150, 10, 100];
let currentZone = d3_find_region_by_point(
    MedievalWorld,
    d3_point_x(playerPos),
    d3_point_y(playerPos),
    d3_point_z(playerPos)
);

if (len(currentZone) > 0) {
    print("Current location: " + currentZone.name);
}

# Calculate world statistics
let totalVol = d3_system_summary(MedievalWorld);
print("Total world volume: " + totalVol);

# Check intersection between castle and a search area
let searchArea = [0, 0, 0, 250, 200, 250];
let intersecting = d3_get_regions_intersecting(
    MedievalWorld,
    searchArea[0], searchArea[1], searchArea[2],
    searchArea[3], searchArea[4], searchArea[5]
);
print("Regions in search area: " + len(intersecting));

print("Spatial program complete!");''')

# ═══════════════════════════════════════════════
#  CHAPTER 16: zzwUI / HwdUI GUI FRAMEWORK
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading_en(doc, 'Chapter 16  zzwUI / HwdUI GUI Framework', level=1)
add_heading_cn(doc, '第十六章  zzwUI / HwdUI GUI框架', level=2)

add_heading_en(doc, '16.1  Overview', level=2)
add_heading_cn(doc, '16.1  概述', level=3)

add_para(doc,
    'zzwUI (also known as HwdUI) is the official GUI widget toolkit for H#. '
    'It provides a complete set of UI components including windows, buttons, labels, text boxes, '
    'checkboxes, radio buttons, sliders, progress bars, images, list boxes, combo boxes, '
    'scroll views, tab controls, and separators. '
    'Starting from version 2.0, HwdUI includes a full CSS-like style engine supporting '
    'selectors, stylesheets, cascade, pseudo-states, and built-in light/dark themes. '
    'The framework is written entirely in H#, achieving full self-hosting within the language.',
    'zzwUI（亦称 HwdUI）是 H# 的官方 GUI 组件工具包。'
    '它提供了一套完整的 UI 组件，包括窗口、按钮、标签、文本框、复选框、单选按钮、'
    '滑块、进度条、图片、列表框、组合框、滚动视图、选项卡和分隔线。'
    '从 2.0 版本起，HwdUI 内置了完整的 CSS-like 样式引擎，支持选择器、样式表、级联、'
    '伪状态和内置的明/暗主题。该框架完全使用 H# 编写，实现了语言内的完全自举。'
)

add_heading_en(doc, '16.2  Core Concepts', level=2)
add_heading_cn(doc, '16.2  核心概念', level=3)

add_para(doc,
    'HwdUI is built around three core abstractions:',
    'HwdUI 围绕三个核心抽象构建：'
)

add_para(doc,
    '1. zzwUI Base Class — All widgets inherit from zzwUI, which provides fundamental '
    'properties (position, size, visibility, enabled state), the parent-child tree structure, '
    'event callback registration, and the CSS-like styling system (class, ID, inline style, '
    'computed style).',
    '1. zzwUI 基类 — 所有组件继承自 zzwUI，它提供了基本属性（位置、大小、可见性、启用状态）、'
    '父子树结构、事件回调注册以及 CSS-like 样式系统（class、ID、inline style、computed style）。'
)

add_para(doc,
    '2. Widgets — Concrete UI components (Window, Button, Panel, etc.) that extend zzwUI '
    'with specific behavior.',
    '2. Widgets（组件）— 具体的 UI 组件（Window、Button、Panel 等），通过扩展 zzwUI 实现特定行为。'
)

add_para(doc,
    '3. CSS Style Engine — A cascading style computation system supporting type selectors, '
    'class selectors, ID selectors, pseudo-class selectors, stylesheets, default styles, '
    'inline styles, and pre-built themes.',
    '3. CSS 样式引擎 — 级联样式计算系统，支持类型选择器、类选择器、ID 选择器、伪类选择器、'
    '样式表、默认样式、内联样式和预构建主题。'
)

add_heading_en(doc, '16.3  Initialization', level=2)
add_heading_cn(doc, '16.3  初始化', level=3)

add_para(doc,
    'Before using any HwdUI widgets, you must initialize the framework by calling hwdui_init(). '
    'This function registers all widget types, resets the ID counter, clears the stylesheet '
    'registry, and resets CSS defaults.',
    '在使用任何 HwdUI 组件之前，必须调用 hwdui_init() 初始化框架。'
    '此函数注册所有组件类型、重置ID计数器、清空样式表注册表和重置CSS默认值。'
)

add_code_block(doc, '''import "bootstrap/hwdui.hto";

hwdui_init();
hwdui_summary();''')

add_heading_en(doc, '16.4  Creating Your First GUI', level=2)
add_heading_cn(doc, '16.4  创建第一个GUI', level=3)

add_para(doc,
    'The following example creates a window with a button that responds to clicks:',
    '以下示例创建一个带有点击响应按钮的窗口：'
)

add_code_block(doc, '''let win = hwdui_create_window("My App", 400, 300);
win.set_pos(100, 100);

let btn = new Button();
btn.init_btn("Click Me");
btn.set_bounds(50, 50, 120, 40);
btn.onClick = fn() {
    print("Button clicked!");
};
win.add_child(btn);

win.show();''')

add_para(doc,
    'Each widget type has its own init method. For Button, use init_btn("text"). '
    'For Label, use init_label("text"). For TextBox, use init_tb("placeholder").',
    '每种组件类型都有自己的初始化方法。Button 使用 init_btn("文本")。'
    'Label 使用 init_label("文本")。TextBox 使用 init_tb("占位文本")。'
)

add_heading_en(doc, '16.5  Layout System', level=2)
add_heading_cn(doc, '16.5  布局系统', level=3)

add_para(doc,
    'HwdUI supports four layout modes, set on Panel widgets via init_panel(layout_type):',
    'HwdUI 支持四种布局模式，通过 init_panel(layout_type) 在 Panel 上设置：'
)

add_para(doc,
    '"absolute" — Child widgets use absolute positioning (x, y, width, height). '
    'This is the default mode and gives full control over widget placement.\n'
    '"vbox" — Children are stacked vertically. Each child\'s y position is computed automatically.\n'
    '"hbox" — Children are arranged horizontally. Each child\'s x position is computed automatically.\n'
    '"grid" — Children are arranged in a grid defined by rows and columns.',
    '"absolute" — 子组件使用绝对定位（x, y, width, height）。这是默认模式，提供完全的位置控制。\n'
    '"vbox" — 子组件垂直堆叠排列。每个子组件的 y 位置自动计算。\n'
    '"hbox" — 子组件水平排列。每个子组件的 x 位置自动计算。\n'
    '"grid" — 子组件在行列网格中排列。'
)

add_code_block(doc, '''let panel = new Panel();
panel.init_panel("vbox");
panel.set_bounds(0, 0, 300, 200);

let btn1 = new Button();
btn1.init_btn("Button 1");
panel.add_child(btn1);

let btn2 = new Button();
btn2.init_btn("Button 2");
panel.add_child(btn2);''')

add_heading_en(doc, '16.6  Event System', level=2)
add_heading_cn(doc, '16.6  事件系统', level=3)

add_para(doc,
    'zzwUI widgets support the following event callbacks, which can be set to H# function values:',
    'zzwUI 组件支持以下事件回调，可设置为 H# 函数值：'
)

add_code_block(doc, '''onClick           onDoubleClick     onKeyDown
onKeyUp           onFocus           onBlur
onMouseEnter      onMouseLeave      onMouseMove
onResize          onMove            onShow
onHide''')

add_para(doc,
    'Event callbacks are set by assigning a function to the corresponding property. '
    'When the event occurs, the widget calls the function automatically.',
    '事件回调通过将函数赋值给对应属性来设置。事件发生时，组件自动调用该函数。'
)

add_code_block(doc, '''let txt = new TextBox();
txt.init_tb("Enter name");
txt.onFocus = fn() {
    print("TextBox focused");
};
txt.onBlur = fn() {
    print("TextBox lost focus, value: " + txt.get_value());
};''')

add_heading_en(doc, '16.7  CSS-like Style Engine', level=2)
add_heading_cn(doc, '16.7  CSS-like 样式引擎', level=3)

add_para(doc,
    'The HwdUI 2.0 CSS engine brings web-like styling to H# GUI programming. '
    'It supports 80+ CSS properties (color, background, border, font, text, spacing, sizing, '
    'display, position, flexbox, grid, shadow, transform, transition, cursor, z-index, opacity), '
    'a full selector system, cascade computation, pseudo-states, and built-in themes. '
    'The entire engine is implemented in pure H# using key-value pair arrays — '
    'no Python builtins are required.',
    'HwdUI 2.0 CSS 引擎为 H# GUI 编程带来了网页般的样式能力。'
    '它支持 80+ CSS 属性（颜色、背景、边框、字体、文本、间距、尺寸、'
    '显示、定位、flexbox、grid、阴影、变换、过渡、光标、z-index、透明度等），'
    '完整的选择器系统、级联计算、伪状态和内置主题。'
    '整个引擎使用键值对数组在纯 H# 中实现——无需任何 Python 内置函数。'
)

add_heading_en(doc, '16.7.1  Selectors', level=2)
add_heading_cn(doc, '16.7.1  选择器', level=3)

add_para(doc,
    'The CSS engine provides a rich set of selector types. '
    'Each selector computes a specificity score used in cascade resolution.',
    'CSS 引擎提供了丰富的选择器类型。每个选择器计算一个用于级联解析的特异性分数。'
)

add_code_block(doc, '''# Type selector — matches all widgets of a given type
hwdui_sel_type("Button")

# Class selector — matches widgets with the given CSS class
hwdui_sel_class("primary")

# ID selector — matches the widget with the exact CSS ID
hwdui_sel_id("submitBtn")

# Compound: type + class
hwdui_sel_type_class("Button", "primary")

# Pseudo-class: type + pseudo state
hwdui_sel_type_pseudo("Button", "hover")

# Universal selector — matches all widgets
hwdui_sel_universal()''')

add_para(doc,
    'Selectors also support parent-based matching (sel_parent_type, sel_parent_class) '
    'for descendant-style rules.',
    '选择器还支持基于父级的匹配（sel_parent_type、sel_parent_class），用于后代样式规则。'
)

add_heading_en(doc, '16.7.2  Stylesheets', level=2)
add_heading_cn(doc, '16.7.2  样式表', level=3)

add_para(doc,
    'A Stylesheet is a collection of StyleRules. Each rule pairs a Selector with '
    'an array of [property, value] declarations.',
    '样式表（Stylesheet）是一组样式规则（StyleRule）。每条规则将选择器与键值对数组配对。'
)

add_code_block(doc, '''let sheet = hwdui_create_stylesheet("my_theme");

sheet.addRule(hwdui_sel_type("Button"), [
    ["background-color", "#e0e0e0"],
    ["color", "#333333"],
    ["border-radius", "4"],
    ["font-size", "14"],
    ["cursor", "pointer"]
]);

sheet.addRule(hwdui_sel_type_pseudo("Button", "hover"), [
    ["background-color", "#d0d0d0"]
]);

sheet.addRule(hwdui_sel_class("danger"), [
    ["background-color", "#ff4444"],
    ["color", "#ffffff"]
]);''')

add_para(doc,
    'Multiple stylesheets can be registered simultaneously. Rules are matched by selector '
    'specificity, and later stylesheets override earlier ones when specificity is equal.',
    '可以同时注册多个样式表。规则按选择器特异性匹配，特异性相同时后注册的样式表覆盖先注册的。'
)

add_heading_en(doc, '16.7.3  Cascade & Computed Style', level=2)
add_heading_cn(doc, '16.7.3  级联与计算样式', level=3)

add_para(doc,
    'The computed style for each widget is resolved through four cascade layers '
    '(lowest to highest priority):',
    '每个组件的计算样式通过四级级联层解析（优先级从低到高）：'
)

add_para(doc,
    'Layer 0: Global defaults — Set via hwdui_set_default_styles(pairs)\n'
    'Layer 1: Stylesheet rules — Matched by selector specificity (higher specificity wins)\n'
    'Layer 2: Legacy styles dict — Set via set_style("key", "value") for backward compatibility\n'
    'Layer 3: Inline styles — Set via setInlineStyle(prop, value), highest priority',
    '第0层：全局默认值 — 通过 hwdui_set_default_styles(pairs) 设置\n'
    '第1层：样式表规则 — 按选择器特异性匹配（高特异性胜出）\n'
    '第2层：传统 styles 字典 — 通过 set_style("key", "value") 设置（向后兼容）\n'
    '第3层：内联样式 — 通过 setInlineStyle(prop, value) 设置，最高优先级'
)

add_code_block(doc, '''# Set global defaults
hwdui_set_default_styles([
    ["font-family", "system-ui"],
    ["font-size", "14"]
]);

# Get computed style (resolves all cascade layers)
let cs = my_button.getComputedStyle();
print("Computed background: " + cs["background-color"]);

# Get single property
let font_size = my_button.getComputedStyleValue("font-size");''')

add_heading_en(doc, '16.7.4  Pseudo States', level=2)
add_heading_cn(doc, '16.7.4  伪状态', level=3)

add_para(doc,
    'Widgets support four CSS pseudo-states: normal, hover, active, focus, and disabled. '
    'Set the current state with setPseudoState(state) and the engine automatically '
    'applies the matching pseudo-class rules.',
    '组件支持四种 CSS 伪状态：normal、hover、active、focus 和 disabled。'
    '通过 setPseudoState(state) 设置当前状态，引擎自动应用匹配的伪类规则。'
)

add_code_block(doc, '''let btn = new Button();
btn.init_btn("Hover Me");

btn.setPseudoState("normal");
let cs_normal = btn.getComputedStyle();

btn.setPseudoState("hover");
let cs_hover = btn.getComputedStyle();

btn.setPseudoState("disabled");
let cs_disabled = btn.getComputedStyle();''')

add_heading_en(doc, '16.7.5  Themes', level=2)
add_heading_cn(doc, '16.7.5  主题', level=3)

add_para(doc,
    'HwdUI ships with two built-in themes. Both are implemented as stylesheets '
    'covering all standard widget types with appropriate pseudo-state rules.',
    'HwdUI 内置两种主题。两者都实现为覆盖所有标准组件类型及适当伪状态规则的样式表。'
)

add_code_block(doc, '''hwdui_clear_all_stylesheets();

# Light theme — clean, modern, with blue accent colors
let light = hwdui_theme_light();

# Dark theme — low-light, high-contrast, with blue focus rings
let dark = hwdui_theme_dark();''')

add_para(doc,
    'Each theme covers: Window, Button (normal/hover/active/disabled), Label, '
    'TextBox (normal/focus/disabled), CheckBox, RadioButton, Slider, ProgressBar, '
    'Panel, ListBox, ComboBox, TabControl, ScrollView, and Separator.',
    '每个主题覆盖：Window、Button（normal/hover/active/disabled）、Label、'
    'TextBox（normal/focus/disabled）、CheckBox、RadioButton、Slider、ProgressBar、'
    'Panel、ListBox、ComboBox、TabControl、ScrollView、Separator。'
)

add_heading_en(doc, '16.7.6  Class & ID Management', level=2)
add_heading_cn(doc, '16.7.6  类与ID管理', level=3)

add_para(doc,
    'Every zzwUI widget manages its own CSS classes (array) and CSS ID (string). '
    'These are independent of the widget\'s id property and are used solely for '
    'selector matching in the CSS engine.',
    '每个 zzwUI 组件管理自己的 CSS 类（数组）和 CSS ID（字符串）。'
    '这些与组件的 id 属性无关，仅用于 CSS 引擎中的选择器匹配。'
)

add_code_block(doc, '''let btn = new Button();
btn.init_btn("Submit");

btn.addClass("primary");
btn.addClass("large");
assert(btn.hasClass("primary") == true, "has class");

btn.toggleClass("active");
btn.removeClass("large");

btn.setCssId("submitBtn");
let css_id = btn.getCssId();

btn.setClasses(["small", "outline"]);''')

add_heading_en(doc, '16.7.7  Inline Styles', level=2)
add_heading_cn(doc, '16.7.7  内联样式', level=3)

add_para(doc,
    'Inline styles have the highest cascade priority and override all stylesheet rules. '
    'They are set per-widget using setInlineStyle or setInlineStyles.',
    '内联样式具有最高级联优先级，覆盖所有样式表规则。'
    '通过 setInlineStyle 或 setInlineStyles 按组件设置。'
)

add_code_block(doc, '''# Set a single inline style
btn.setInlineStyle("background-color", "#ff0000");
btn.setInlineStyle("border-radius", "8");

# Set multiple inline styles at once
btn.setInlineStyles([
    ["background-color", "#00ff00"],
    ["color", "#ffffff"],
    ["font-size", "16"]
]);

# Remove an inline style — falls back to stylesheet
btn.removeInlineStyle("background-color");

# Read an inline style
let bg = btn.getInlineStyle("background-color");''')

add_heading_en(doc, '16.8  Complete Widget Reference', level=2)
add_heading_cn(doc, '16.8  完整组件参考', level=3)

add_para(doc,
    'The following table lists all 16 HwdUI widget types, their init methods, and key properties:',
    '下表列出了所有 16 种 HwdUI 组件类型、初始化方法和关键属性：'
)

widget_table = [
    ('zzwUI', 'init("id")', 'Base class for all widgets. Provides position, size, visibility, '
     'enabled state, parent-child tree, event callbacks, CSS class/ID/inline style management, '
     'and computed style resolution.',
     '所有组件的基类。提供位置、大小、可见性、启用状态、父子树、事件回调、CSS class/ID/内联样式管理和计算样式解析。'),
    ('Window', 'hwdui_create_window("title", w, h)', 'Top-level application window with title bar, '
     'minimize/maximize/restore/close. Contains a client area for child widgets.',
     '顶层应用窗口，带标题栏、最小化/最大化/还原/关闭。包含用于子组件的客户区。'),
    ('Panel', 'init_panel("layout_type")', 'Container widget. Supports absolute, vbox, hbox, and grid '
     'layout modes. Best for organizing child widgets.',
     '容器组件。支持 absolute、vbox、hbox、grid 四种布局模式。最适合组织子组件。'),
    ('Button', 'init_btn("text")', 'Clickable button with text label. Supports enabled/disabled states '
     'and onClick event. CSS pseudo-states: hover, active, disabled.',
     '可点击按钮，带文本标签。支持启用/禁用状态和 onClick 事件。CSS 伪状态：hover、active、disabled。'),
    ('Label', 'init_label("text")', 'Non-interactive text display widget. Used for titles, '
     'descriptions, and static text.',
     '非交互式文本显示组件。用于标题、描述和静态文本。'),
    ('TextBox', 'init_tb("placeholder")', 'Single-line text input field. Supports focus/blur events, '
     'text selection, and placeholder text. CSS pseudo-states: focus, disabled.',
     '单行文本输入框。支持 focus/blur 事件、文本选择和占位文本。CSS 伪状态：focus、disabled。'),
    ('CheckBox', 'init_cb("label")', 'Toggleable checkbox with label. Supports checked/unchecked state '
     'and onChange event.',
     '可切换复选框，带标签。支持选中/未选中状态和 onChange 事件。'),
    ('RadioButton', 'init_rb("label")', 'Radio button for mutually exclusive selection within a group.',
     '单选按钮，用于组内互斥选择。'),
    ('Slider', 'init_slider(min, max, value)', 'Horizontal slider for selecting a numeric value '
     'within a range. Supports onChange event.',
     '水平滑块，用于在范围内选择数值。支持 onChange 事件。'),
    ('ProgressBar', 'init_pb(min, max, value)', 'Horizontal progress bar showing completion percentage.',
     '水平进度条，显示完成百分比。'),
    ('Image', 'init_image("path")', 'Displays an image from a file path. Supports scaling and '
     'fit modes.',
     '从文件路径显示图片。支持缩放和适配模式。'),
    ('ListBox', 'init_lb()', 'Scrollable list of items. Supports selection, multi-select, '
     'add_item/remove_item.',
     '可滚动列表。支持选择、多选、add_item/remove_item。'),
    ('ComboBox', 'init_cb("placeholder")', 'Dropdown selection widget combining a text field '
     'with a popup list. Supports add_item and selection events.',
     '下拉选择组件，结合文本框和弹出列表。支持 add_item 和选择事件。'),
    ('ScrollView', 'init_sv(content_w, content_h)', 'Scrollable container for content larger '
     'than the visible area. Supports horizontal and vertical scrolling.',
     '可滚动容器，用于超出可见区域的内容。支持水平和垂直滚动。'),
    ('TabControl', 'init_tc()', 'Tabbed interface widget. add_tab("name", content_widget) adds tabs.',
     '选项卡界面组件。add_tab("name", content_widget) 添加选项卡。'),
    ('Separator', 'init_sep("horizontal")', 'Visual divider line. Supports horizontal or vertical '
     'orientation.',
     '视觉分隔线。支持水平或垂直方向。'),
]

for name, init_sig, en_desc, cn_desc in widget_table:
    add_heading_en(doc, name, level=3)
    add_para(doc, en_desc, cn_desc)
    add_code_block(doc, init_sig)

add_heading_en(doc, '16.9  Complete GUI Application Example', level=2)
add_heading_cn(doc, '16.9  完整GUI应用示例', level=3)

add_para(doc,
    'The following example demonstrates a complete HwdUI application using the CSS style engine, '
    'class-based styling, pseudo-states, and event handling:',
    '以下示例演示了一个使用 CSS 样式引擎、类样式、伪状态和事件处理的完整 HwdUI 应用：'
)

add_code_block(doc, '''import "bootstrap/hwdui.hto";

hwdui_init();

# Apply the built-in light theme
hwdui_theme_light();

# Set global defaults
hwdui_set_default_styles([
    ["font-family", "sans-serif"],
    ["font-size", "13"]
]);

# Create main window
let main_win = hwdui_create_window("My HwdUI App", 500, 400);
main_win.set_pos(200, 100);

# Create a panel with vbox layout
let main_panel = new Panel();
main_panel.init_panel("vbox");
main_panel.set_bounds(10, 40, 480, 350);

# Title label with custom inline style
let title_lbl = new Label();
title_lbl.init_label("Welcome to HwdUI 2.0");
title_lbl.set_bounds(0, 0, 480, 30);
title_lbl.setInlineStyles([
    ["font-size", "20"],
    ["font-weight", "bold"],
    ["text-align", "center"]
]);
main_panel.add_child(title_lbl);

# Spacer
let spacer = new Panel();
spacer.init_panel("absolute");
spacer.set_bounds(0, 0, 480, 10);
main_panel.add_child(spacer);

# Button with class-based styling
let primary_btn = new Button();
primary_btn.init_btn("Primary Action");
primary_btn.set_bounds(0, 0, 200, 40);
primary_btn.addClass("primary");

primary_btn.onClick = fn() {
    print("Primary button clicked!");
    primary_btn.setInlineStyle("background-color", "#00aa00");
};
main_panel.add_child(primary_btn);

# Button with danger class
let danger_btn = new Button();
danger_btn.init_btn("Danger Action");
danger_btn.set_bounds(0, 0, 200, 40);
danger_btn.addClass("danger");
danger_btn.setInlineStyle("background-color", "#cc0000");

danger_btn.onMouseEnter = fn() {
    danger_btn.setPseudoState("hover");
};
danger_btn.onMouseLeave = fn() {
    danger_btn.setPseudoState("normal");
};
main_panel.add_child(danger_btn);

# Text input
let name_input = new TextBox();
name_input.init_tb("Enter your name...");
name_input.set_bounds(0, 0, 280, 30);
main_panel.add_child(name_input);

# Status label
let status_lbl = new Label();
status_lbl.init_label("Status: Ready");
status_lbl.set_bounds(0, 0, 480, 20);
status_lbl.setInlineStyle("color", "#888888");
main_panel.add_child(status_lbl);

# Add panel to window and show
main_win.add_child(main_panel);
main_win.show();

# Print computed styles for debugging
print("\\n=== Computed Styles ===");
let cs = primary_btn.getComputedStyle();
print("Background: " + cs["background-color"]);
print("Font size: " + cs["font-size"]);''')

add_heading_en(doc, '16.10  Supported CSS Properties', level=2)
add_heading_cn(doc, '16.10  支持的CSS属性', level=3)

add_para(doc,
    'The HwdUI CSS engine supports over 80 CSS-like properties, organized by category:',
    'HwdUI CSS 引擎支持超过 80 个 CSS-like 属性，按类别组织如下：'
)

css_categories = [
    ('Color & Background / 颜色与背景',
     'color, background-color, background, background-image, '
     'background-size, background-position, background-repeat, opacity'),
    ('Border / 边框',
     'border, border-color, border-width, border-style, border-radius, '
     'border-top/right/bottom/left-width, border-top/right/bottom/left-color, '
     'border-top/right/bottom/left-style, '
     'border-top-left/top-right/bottom-left/bottom-right-radius'),
    ('Font & Text / 字体与文本',
     'font-family, font-size, font-weight, font-style, '
     'text-align, text-decoration, text-transform, text-shadow, '
     'letter-spacing, line-height, word-spacing, text-indent, white-space'),
    ('Spacing / 间距',
     'margin, margin-top/right/bottom/left, '
     'padding, padding-top/right/bottom/left'),
    ('Sizing / 尺寸',
     'width, height, min-width, min-height, max-width, max-height'),
    ('Display & Position / 显示与定位',
     'display, position, top, right, bottom, left, '
     'visibility, overflow, overflow-x, overflow-y, z-index, cursor'),
    ('Flexbox',
     'flex-direction, flex-wrap, justify-content, align-items, '
     'align-content, flex-grow, flex-shrink, flex-basis, align-self, order, gap'),
    ('Grid',
     'grid-template-columns, grid-template-rows, grid-column, grid-row, grid-gap'),
    ('Effects / 效果',
     'box-shadow, transform, transform-origin, '
     'transition, transition-duration, transition-property, '
     'transition-timing-function, transition-delay'),
]

for cat_name, props in css_categories:
    add_para(doc, cat_name, cat_name)
    add_code_block(doc, props)

add_heading_en(doc, '16.11  API Reference', level=2)
add_heading_cn(doc, '16.11  API参考', level=3)

add_para(doc,
    'Core functions of the HwdUI CSS engine:',
    'HwdUI CSS 引擎的核心函数：'
)

add_code_block(doc, '''# Framework Initialization
hwdui_init()                — Initialize the framework
hwdui_summary()             — Print framework summary

# Selector Factory Functions
hwdui_sel_type(t)             — Type selector
hwdui_sel_class(c)            — Class selector
hwdui_sel_id(i)               — ID selector
hwdui_sel_type_class(t, c)    — Type + class compound selector
hwdui_sel_type_pseudo(t, p)   — Type + pseudo-class selector
hwdui_sel_universal()         — Universal selector (*)

# Stylesheet Management
hwdui_create_stylesheet(name) — Create and register a stylesheet
hwdui_remove_stylesheet(s)    — Remove a stylesheet
hwdui_clear_all_stylesheets() — Remove all stylesheets

# Style Computation
hwdui_compute_style(widget)   — Compute final cascaded style
hwdui_get_computed_prop(w, p) — Get single computed property
hwdui_print_computed_style(w) — Print computed style (debug)
hwdui_get_computed_margin(w)  — Get computed margin dict
hwdui_get_computed_padding(w) — Get computed padding dict

# Themes
hwdui_theme_light()           — Apply light theme
hwdui_theme_dark()            — Apply dark theme

# Global Defaults
hwdui_set_default_styles(pairs) — Set global CSS defaults

# Per-Widget CSS Methods (on zzwUI)
widget.addClass(name)           — Add a CSS class
widget.removeClass(name)        — Remove a CSS class
widget.hasClass(name)           — Check for CSS class
widget.toggleClass(name)        — Toggle CSS class
widget.setCssId(id)             — Set CSS ID
widget.getCssId()               — Get CSS ID
widget.setPseudoState(state)    — Set pseudo-state
widget.getPseudoState()         — Get pseudo-state
widget.getComputedStyle()       — Get cascaded computed style
widget.getComputedStyleValue(p) — Get single computed property
widget.setInlineStyle(prop, v)  — Set inline style
widget.getInlineStyle(prop)     — Get inline style
widget.removeInlineStyle(prop)  — Remove inline style
widget.setInlineStyles(pairs)   — Set multiple inline styles

# Legacy Style API
widget.set_style(key, value)    — Set legacy style
widget.get_style(key)           — Get legacy style
widget.apply_styles(dict)       — Apply legacy style dict''')

add_note(doc,
    'The CSS engine stores all style data as arrays of [key, value] pairs rather than '
    'H# dictionaries. This avoids the limitation that H# cannot check dict key existence. '
    'Lookup is performed by linear iteration via the hwdui_pair_get helper.',
    'CSS 引擎将所有样式数据存储为 [key, value] 键值对数组而非 H# 字典。'
    '这避免了 H# 无法检查字典键存在性的限制。查找通过 hwdui_pair_get 辅助函数线性迭代完成。'
)

# ═══════════════════════════════════════════════
#  APPENDIX A: QUICK REFERENCE
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading_en(doc, 'Appendix A  Quick Reference', level=1)
add_heading_cn(doc, '附录A  快速参考', level=2)

add_heading_en(doc, 'A.1  Keywords', level=2)
add_code_block(doc, '''let fn return while if else for in print import
class new extends private static interface implements super
is as module concept coro asm ptr true false and or not
continue break nullptr auto try catch throw''')

add_heading_en(doc, 'A.2  Operators', level=2)
add_code_block(doc, '''+   -   *   /       Arithmetic / 算术
==  !=  >   <  >=  <= Comparison / 比较
and or not           Logical / 逻辑
&   |   ^   ~        Bitwise / 位运算
<<  >>               Shift / 移位
=                    Assignment / 赋值''')

add_heading_en(doc, 'A.3  Literals', level=2)
add_code_block(doc, '''42, -7, 3.14         Numbers / 数字
"hello"              Strings / 字符串
true, false          Booleans / 布尔值
nullptr              Null value / 空值
[1, 2, 3]            Array literal / 数组字面量
[["k","v"]]          Dict literal / 字典字面量''')

add_heading_en(doc, 'A.4  Control Flow Syntax', level=2)
add_code_block(doc, '''if (cond) { ... }
if (cond) { ... } else { ... }
if (cond) { ... } else if (cond2) { ... } else { ... }
while (cond) { ... }
for (let x in arr) { ... }
for (let k, v in pairs) { ... }
break;
continue;
return expr;''')

add_heading_en(doc, 'A.5  Function Syntax', level=2)
add_code_block(doc, '''fn name(p1, p2) {
    ...
    return value;
}

# Nested function / 嵌套函数
fn outer(x) {
    fn inner(y) {
        return x + y;
    }
    return inner;
}''')

add_heading_en(doc, 'A.6  Class Syntax', level=2)
add_code_block(doc, '''class Name {
    let field = value;
    private let hidden = value;
    static fn method() { ... }
    fn method() { ... }
}

class Child extends Parent implements Iface { ... }
interface Iface { fn method(); }
concept C { fn helper() { ... } }''')

add_heading_en(doc, 'A.7  Common Built-in Functions', level=2)
add_code_block(doc, '''len(x)          push(a, v)      pop(a)
read_file(p)    write_file(p,s) input(p)
int(x)          str(x)          ord(c)      chr(n)
substring(s,b,e) time_now()
http_get(u)     http_post(u,d)  json_stringify(o)  json_parse(s)''')

# ═══════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           'HSharp_Programming_Tutorial.docx')
doc.save(output_path)
print(f'Tutorial saved to: {output_path}')