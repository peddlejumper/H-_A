#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
《H# 从入门到精通》150 页完整版 — Tech Innovation 主题
主框架:样式定义 + 封面 + 前言 + 目录 + 集成 35 章 + 附录
"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ==================== Tech Innovation 主题色 ====================
C_PRIMARY  = RGBColor(0x00, 0x66, 0xFF)   # 电蓝 — 主标题/强调
C_ACCENT   = RGBColor(0x00, 0x99, 0xB3)   # 霓虹青(调暗) — 副强调
C_DARK     = RGBColor(0x1E, 0x1E, 0x1E)   # 深灰 — 章标题
C_TEXT     = RGBColor(0x33, 0x33, 0x33)   # 正文
C_CODE_BG  = 'F0F4F8'                      # 代码块背景(淡蓝灰)
C_OUTPUT   = RGBColor(0x00, 0x70, 0x2A)   # 输出绿
C_NOTE     = RGBColor(0xCC, 0x66, 0x00)   # 提示橙
C_WARN     = RGBColor(0xCC, 0x00, 0x00)   # 警告红
C_MUTED    = RGBColor(0x88, 0x88, 0x88)   # 次要灰
C_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)

FONT_CN = '微软雅黑'
FONT_EN = 'Consolas'

doc = Document()

# ==================== 全局样式 ====================
normal = doc.styles['Normal']
normal.font.name = FONT_EN
normal.font.size = Pt(10.5)
normal.font.color.rgb = C_TEXT
normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
normal.paragraph_format.line_spacing = 1.35
normal.paragraph_format.space_after = Pt(4)

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# 标题样式
for i, sz in [(1, 26), (2, 20), (3, 14), (4, 12)]:
    hs = doc.styles[f'Heading {i}']
    hs.font.name = FONT_CN
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
    hs.font.size = Pt(sz)
    hs.font.bold = True
    hs.font.color.rgb = C_PRIMARY if i <= 2 else C_DARK

# ==================== 辅助函数 ====================

def _set_run_font(run, en=FONT_EN, cn=FONT_CN, size=None, color=None, bold=None, italic=None):
    run.font.name = en
    run.element.rPr.rFonts.set(qn('w:eastAsia'), cn)
    if size:  run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if bold is not None:  run.font.bold = bold
    if italic is not None: run.font.italic = italic

def _set_para_shading(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)

def _set_para_border(paragraph, color='0066FF', size='6', position='bottom'):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    edge = OxmlElement(f'w:{position}')
    edge.set(qn('w:val'), 'single')
    edge.set(qn('w:sz'), size)
    edge.set(qn('w:space'), '4')
    edge.set(qn('w:color'), color)
    pbdr.append(edge)
    pPr.append(pbdr)


class H:
    """书籍内容辅助类 — 各篇章模块通过此类向 doc 添加内容"""

    @staticmethod
    def h1(text):
        """篇标题:居中大字号电蓝,前分页(使用 Heading 1 样式便于导航)"""
        doc.add_page_break()
        for _ in range(5):
            doc.add_paragraph()
        p = doc.add_paragraph(style='Heading 1')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        _set_run_font(run, cn=FONT_CN, size=28, color=C_PRIMARY, bold=True)
        doc.add_paragraph()

    @staticmethod
    def h2(text):
        """章标题:深灰粗体 + 电蓝底线"""
        doc.add_paragraph()
        p = doc.add_paragraph(style='Heading 2')
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(text)
        _set_run_font(run, cn=FONT_CN, size=20, color=C_DARK, bold=True)
        _set_para_border(p, color='0066FF', size='8', position='bottom')

    @staticmethod
    def h3(text):
        """节标题"""
        p = doc.add_paragraph(style='Heading 3')
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        _set_run_font(run, cn=FONT_CN, size=14, color=C_DARK, bold=True)

    @staticmethod
    def h4(text):
        """子标题"""
        p = doc.add_paragraph(style='Heading 4')
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        _set_run_font(run, cn=FONT_CN, size=12, color=C_ACCENT, bold=True)

    @staticmethod
    def para(text):
        """正文段落"""
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text)
        _set_run_font(run, cn=FONT_CN, size=10.5, color=C_TEXT)

    @staticmethod
    def code(text):
        """代码块:Consolas 9pt,淡蓝灰背景"""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.3)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        _set_para_shading(p, C_CODE_BG)
        lines = text.split('\n')
        for i, line in enumerate(lines):
            run = p.add_run(line)
            _set_run_font(run, en=FONT_EN, cn=FONT_EN, size=9, color=C_DARK)
            if i < len(lines) - 1:
                run.add_break()

    @staticmethod
    def output(text):
        """程序输出:Consolas 9pt 绿色"""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        prefix = p.add_run('输出:\n')
        _set_run_font(prefix, en=FONT_EN, cn=FONT_CN, size=9, color=C_MUTED, italic=True)
        lines = text.split('\n')
        for i, line in enumerate(lines):
            run = p.add_run(line)
            _set_run_font(run, en=FONT_EN, cn=FONT_EN, size=9, color=C_OUTPUT)
            if i < len(lines) - 1:
                run.add_break()

    @staticmethod
    def note(text):
        """提示框:橙色斜体"""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(6)
        _set_para_shading(p, 'FFF8E7')
        icon = p.add_run('提示  ')
        _set_run_font(icon, cn=FONT_CN, size=10, color=C_NOTE, bold=True)
        run = p.add_run(text)
        _set_run_font(run, cn=FONT_CN, size=10, color=C_NOTE, italic=True)

    @staticmethod
    def warning(text):
        """警告框:红色粗体"""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(6)
        _set_para_shading(p, 'FFF0F0')
        icon = p.add_run('注意  ')
        _set_run_font(icon, cn=FONT_CN, size=10, color=C_WARN, bold=True)
        run = p.add_run(text)
        _set_run_font(run, cn=FONT_CN, size=10, color=C_WARN, bold=True)

    @staticmethod
    def bullet(text):
        """项目符号"""
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        _set_run_font(run, cn=FONT_CN, size=10.5, color=C_TEXT)

    @staticmethod
    def number(text):
        """编号列表项"""
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        _set_run_font(run, cn=FONT_CN, size=10.5, color=C_TEXT)

    @staticmethod
    def page_break():
        doc.add_page_break()

    @staticmethod
    def blank():
        doc.add_paragraph()


# ==================== 封面 ====================
for _ in range(6):
    doc.add_paragraph()

# 顶部装饰线
_deco = doc.add_paragraph()
_deco.alignment = WD_ALIGN_PARAGRAPH.CENTER
_set_para_border(_deco, color='0066FF', size='12', position='bottom')

for _ in range(2):
    doc.add_paragraph()

_title = doc.add_paragraph()
_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = _title.add_run('H# 从入门到精通')
_set_run_font(run, cn=FONT_CN, size=44, color=C_PRIMARY, bold=True)

_sub = doc.add_paragraph()
_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = _sub.add_run('H# Programming Language: From Beginner to Expert')
_set_run_font(run, en=FONT_EN, cn=FONT_CN, size=14, color=C_ACCENT, italic=True)

for _ in range(2):
    doc.add_paragraph()

_desc = doc.add_paragraph()
_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = _desc.add_run('基于 H# v0.4.1  |  35 章 + 附录  |  全部代码示例经实机测试')
_set_run_font(run, cn=FONT_CN, size=12, color=C_MUTED)

for _ in range(8):
    doc.add_paragraph()

# 底部装饰线
_deco2 = doc.add_paragraph()
_deco2.alignment = WD_ALIGN_PARAGRAPH.CENTER
_set_para_border(_deco2, color='00B3B3', size='12', position='top')

for _ in range(2):
    doc.add_paragraph()

_ed = doc.add_paragraph()
_ed.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = _ed.add_run('2026 年 · 技术创新版')
_set_run_font(run, cn=FONT_CN, size=11, color=C_MUTED)

doc.add_page_break()

# ==================== 前言 ====================
_pref = doc.add_paragraph()
run = _pref.add_run('前  言')
_set_run_font(run, cn=FONT_CN, size=22, color=C_PRIMARY, bold=True)
_pref.alignment = WD_ALIGN_PARAGRAPH.CENTER
_set_para_border(_pref, color='0066FF', size='8', position='bottom')
doc.add_paragraph()

H.para('编程语言的诞生,往往源于对「简洁」与「表达力」的双重追求。H# 正是这样一门语言——它以 Python 为根基,汲取了现代脚本语言的直观语法,同时融合了静态类型系统的部分思想与并发编程模型,旨在为学习者提供一条从「入门」到「精通」的平滑路径。')
H.para('H# 并非一门只存在于纸面上的语言。它拥有三套完整的运行时实现:Python 树遍历解释器用于开发与调试,Python 字节码虚拟机提供更高的执行效率,Kotlin HVM 栈式虚拟机则面向生产部署。更令人兴奋的是,H# 的工具链(tokenizer、parser、compiler、interpreter)可以用 H# 自身来实现——这就是 bootstrap 自举,一门语言成熟度的重要标志。')
H.para('本书的目标,是带领读者从零开始,系统、深入地掌握 H# 语言的方方面面。无论是第一次接触编程的新手,还是已有其他语言经验的开发者,都能在本书中找到适合自己的学习节奏。')

_h3 = doc.add_paragraph(style='Heading 3')
run = _h3.add_run('本书读者对象')
_set_run_font(run, cn=FONT_CN, size=14, color=C_DARK, bold=True)
H.bullet('编程初学者:希望以一门语法友好、生态完整语言入门')
H.bullet('Python 开发者:想了解一门 Python 衍生语言的设计差异与扩展特性')
H.bullet('语言爱好者与编译原理学习者:对自举、字节码 VM、多运行时架构感兴趣')
H.bullet('后端工程师:关注并发编程(async/await、Channel)与高效部署')

_h3 = doc.add_paragraph(style='Heading 3')
run = _h3.add_run('本书如何组织')
_set_run_font(run, cn=FONT_CN, size=14, color=C_DARK, bold=True)
H.para('全书共分为八大篇、35 章,并附有 4 个速查附录,内容由浅入深、循序渐进:')
H.bullet('第一篇 基础入门(第 1–4 章):语言简介、环境搭建、语法元素、变量与赋值')
H.bullet('第二篇 核心语法(第 5–9 章):数据类型、运算符、控制流、字符串、输入输出')
H.bullet('第三篇 函数式编程(第 10–13 章):函数、作用域与闭包、Lambda、函数式实践')
H.bullet('第四篇 面向对象(第 14–17 章):类与对象、继承多态、封装、OOP 实战')
H.bullet('第五篇 数据结构与算法(第 18–22 章):列表、字典、字符串进阶、排序查找、数据结构')
H.bullet('第六篇 高级特性(第 23–27 章):异常处理、Union 类型、模块、接口、错误传播')
H.bullet('第七篇 并发编程(第 28–31 章):async/await、Channel、结构化并发、并发实战')
H.bullet('第八篇 工具链与生态(第 32–35 章):标准库、字节码 VM、Bootstrap 自举、打包发布')

_h3 = doc.add_paragraph(style='Heading 3')
run = _h3.add_run('如何阅读本书')
_set_run_font(run, cn=FONT_CN, size=14, color=C_DARK, bold=True)
H.para('初学者建议按顺序从第一篇阅读,每章的代码示例都应在本地动手运行一遍。有编程经验的读者可以快速浏览前两篇,将精力集中在函数式编程、面向对象与并发编程等核心篇章。每章末尾附有「本章小结」与「练习题」,建议独立完成以巩固所学。')
H.para('本书所有代码示例均经过 interpreter.py 实机测试,输出结果与书中展示完全一致。遇到运行结果不一致时,请确认所使用的 H# 版本为 v0.4.1。')

_h3 = doc.add_paragraph(style='Heading 3')
run = _h3.add_run('勘误与反馈')
_set_run_font(run, cn=FONT_CN, size=14, color=C_DARK, bold=True)
H.para('由于 H# 语言仍在快速演进,部分特性在不同运行时(Python 解释器 / Kotlin HVM)上的支持程度可能有所差异。书中已对这类差异作出标注。如发现错误或有改进建议,欢迎反馈。')

doc.add_paragraph()
_sign = doc.add_paragraph()
_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = _sign.add_run('作  者')
_set_run_font(run, cn=FONT_CN, size=11, color=C_MUTED)
_sign2 = doc.add_paragraph()
_sign2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = _sign2.add_run('2026 年 7 月')
_set_run_font(run, cn=FONT_CN, size=11, color=C_MUTED)

doc.add_page_break()

# ==================== 目录 ====================
_toc = doc.add_paragraph()
run = _toc.add_run('目  录')
_set_run_font(run, cn=FONT_CN, size=22, color=C_PRIMARY, bold=True)
_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
_set_para_border(_toc, color='0066FF', size='8', position='bottom')
doc.add_paragraph()

TOC_ENTRIES = [
    ('第一篇 基础入门', True),
    ('第 1 章  H# 语言简介', False),
    ('第 2 章  环境搭建', False),
    ('第 3 章  基本语法元素', False),
    ('第 4 章  变量与赋值', False),
    ('第二篇 核心语法', True),
    ('第 5 章  数据类型', False),
    ('第 6 章  运算符', False),
    ('第 7 章  控制流', False),
    ('第 8 章  字符串详解', False),
    ('第 9 章  输入输出', False),
    ('第三篇 函数式编程', True),
    ('第 10 章  函数定义与调用', False),
    ('第 11 章  作用域与闭包', False),
    ('第 12 章  Lambda 表达式', False),
    ('第 13 章  函数式编程实践', False),
    ('第四篇 面向对象', True),
    ('第 14 章  类与对象', False),
    ('第 15 章  继承与多态', False),
    ('第 16 章  封装与访问控制', False),
    ('第 17 章  OOP 实战', False),
    ('第五篇 数据结构与算法', True),
    ('第 18 章  列表深入', False),
    ('第 19 章  字典深入', False),
    ('第 20 章  字符串处理进阶', False),
    ('第 21 章  常用算法', False),
    ('第 22 章  数据结构实战', False),
    ('第六篇 高级特性', True),
    ('第 23 章  异常处理', False),
    ('第 24 章  Union 类型', False),
    ('第 25 章  模块与导入', False),
    ('第 26 章  概念与接口', False),
    ('第 27 章  错误传播', False),
    ('第七篇 并发编程', True),
    ('第 28 章  async/await', False),
    ('第 29 章  Channel 通道', False),
    ('第 30 章  并行与结构化并发', False),
    ('第 31 章  并发实战', False),
    ('第八篇 工具链与生态', True),
    ('第 32 章  标准库', False),
    ('第 33 章  字节码 VM', False),
    ('第 34 章  Bootstrap 自举', False),
    ('第 35 章  打包与发布', False),
    ('附  录', True),
    ('附录 A  运算符优先级表', False),
    ('附录 B  内置函数速查', False),
    ('附录 C  关键字列表', False),
    ('附录 D  H# v0.4.1 特性总结', False),
]

for text, is_part in TOC_ENTRIES:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if is_part:
        run = p.add_run(text)
        _set_run_font(run, cn=FONT_CN, size=12, color=C_PRIMARY, bold=True)
        p.paragraph_format.space_before = Pt(8)
    else:
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(text)
        _set_run_font(run, cn=FONT_CN, size=10.5, color=C_TEXT)

doc.add_page_break()

# ==================== 集成各篇章 ====================
import part1, part2, part3, part4, part5

PARTS = [
    ('第一篇 + 第二篇', part1),
    ('第三篇 + 第四篇', part2),
    ('第五篇', part3),
    ('第六篇 + 第七篇', part4),
    ('第八篇 + 附录', part5),
]

for label, mod in PARTS:
    print(f'正在生成: {label} ...')
    mod.add_content(doc, H)
    print(f'  {label} 完成')

# ==================== 尾页 ====================
doc.add_page_break()
for _ in range(8):
    doc.add_paragraph()

_end = doc.add_paragraph()
_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = _end.add_run('— 全书完 —')
_set_run_font(run, cn=FONT_CN, size=28, color=C_PRIMARY, bold=True)

doc.add_paragraph()
_info = doc.add_paragraph()
_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = _info.add_run('基于 H# v0.4.1  |  35 章 + 4 附录  |  全部代码示例经实机测试通过')
_set_run_font(run, cn=FONT_CN, size=11, color=C_MUTED)

doc.add_paragraph()
_run2 = doc.add_paragraph()
_run2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = _run2.add_run('Python 解释器 + Kotlin HVM 双运行时验证  |  Tech Innovation 主题')
_set_run_font(run, cn=FONT_CN, size=10, color=C_MUTED, italic=True)

# ==================== 保存 ====================
output_path = '/Users/peddlejumper/H#/v0.4/H#从入门到精通.docx'
doc.save(output_path)
size_kb = os.path.getsize(output_path) / 1024
print(f'\n文档已生成: {output_path}')
print(f'文件大小: {size_kb:.1f} KB')
