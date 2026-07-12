#!/usr/bin/env python3
"""
Generate a bilingual (Chinese + English) H# tutorial Word document.

Usage:
  python3 scripts/generate_hsharp_tutorial.py

This script requires `python-docx` (install with `pip install python-docx`).
"""
import os
from docx import Document
from docx.shared import Pt


def add_paragraph_cn_en(doc, cn, en):
    p_cn = doc.add_paragraph()
    r_cn = p_cn.add_run(cn)
    r_cn.font.size = Pt(11)

    p_en = doc.add_paragraph()
    r_en = p_en.add_run(en)
    r_en.italic = True
    r_en.font.size = Pt(10)


def add_code_block(doc, code):
    p = doc.add_paragraph()
    r = p.add_run(code)
    try:
        r.font.name = 'Courier New'
    except Exception:
        pass
    r.font.size = Pt(9)


def build_doc(outpath):
    doc = Document()
    doc.core_properties.title = 'H# 程序设计教程 / H# Programming Tutorial'

    doc.add_heading('H# 程序设计教程', level=0)
    doc.add_paragraph('中文 / English bilingual tutorial for H# (v0.4)')

    # Sections: concise bilingual content extracted / translated from repo docs
    sections = [
        {
            'title_cn': '快速开始',
            'title_en': 'Quick Start',
            'body_cn': (
                '先决条件：Python 3.x。\n'
                '运行方式：`python3 hsharp.py path/to/file.hto`，或在 REPL 中交互式运行。'
            ),
            'body_en': (
                'Prerequisites: Python 3.x.\n'
                'Run: `python3 hsharp.py path/to/file.hto` or use the REPL for interactive experimentation.'
            ),
            'code': """let x = 1;
print(x);
""",
        },
        {
            'title_cn': '基本语法与数据类型',
            'title_en': 'Basic Syntax and Data Types',
            'body_cn': (
                '变量声明：`let name = expr;`。支持数字、字符串、布尔、null、数组和字典字面量。函数用 `fn` 定义。'
            ),
            'body_en': (
                'Variables: `let name = expr;`. Literals: numbers, strings, booleans, null, arrays and dictionaries. '
                'Functions are declared with `fn`.'
            ),
            'code': """fn add(a, b) {
    return a + b;
}
""",
        },
        {
            'title_cn': '控制流与集合',
            'title_en': 'Control Flow and Collections',
            'body_cn': (
                '条件：`if (cond) { ... } else { ... }`。循环：`while` 与 `for` 迭代。索引访问和成员访问如 `a[0]`、`obj.field`。'
            ),
            'body_en': (
                'Conditionals: `if (cond) { ... } else { ... }`. Loops: `while` and `for` iteration. '
                'Indexing and member access use `a[0]` and `obj.field`.'
            ),
            'code': """if (x > 0) {
    print("positive");
} else {
    print("non-positive");
}
""",
        },
        {
            'title_cn': '函数、Lambda 与闭包',
            'title_en': 'Functions, Lambdas and Closures',
            'body_cn': (
                '函数是一等公民，支持匿名函数（`fn(...) { ... }`）与闭包。示例展示如何捕获外部变量。'
            ),
            'body_en': (
                'Functions are first-class values; anonymous functions use `fn(...) { ... }`. Closures capture surrounding variables.'
            ),
            'code': """let x = 42;
let f = fn() { print(x); };
f();
""",
        },
        {
            'title_cn': '异常处理',
            'title_en': 'Exception Handling',
            'body_cn': (
                '使用 `try { ... } catch (e) { ... }` 进行异常捕获，内部会生成相应的字节码（例如 `RAISE` / `POP_EXCEPT`）。'
            ),
            'body_en': (
                'Use `try { ... } catch (e) { ... }` to catch exceptions; the compiler emits bytecode like `RAISE` and `POP_EXCEPT` to implement it.'
            ),
            'code': """try {
    throw "oops";
} catch (e) {
    print(e);
}
""",
        },
        {
            'title_cn': '类与面向对象',
            'title_en': 'Classes and Object-Oriented Programming',
            'body_cn': (
                '支持类、方法与字段；运行时用字典实现类/实例结构，支持方法调用与继承（概览）。'
            ),
            'body_en': (
                'Classes, methods and fields are supported. At runtime classes/instances are represented as dictionaries; method calls and inheritance are supported (overview).' 
            ),
            'code': """class Point {
    let x = 0;
    let y = 0;
    fn init(self, x, y) { self.x = x; self.y = y; }
}
""",
        },
        {
            'title_cn': '标准库与内建函数',
            'title_en': 'Standard Library and Builtins',
            'body_cn': (
                '常用函数：`len`, `push`, `pop`, `read_file`, `write_file`。Python 宿主还提供文件与时间等系统接口。'
            ),
            'body_en': (
                'Common builtins: `len`, `push`, `pop`, `read_file`, `write_file`. The Python host exposes file and time APIs.'
            ),
            'code': """let s = read_file("test.txt");
print(s);
""",
        },
        {
            'title_cn': '字节码与 VM 指令集',
            'title_en': 'Bytecode and VM Instruction Set',
            'body_cn': (
                '常见指令包括 `LOAD_CONST`, `CALL_FUNCTION`, `RETURN_VALUE`, `JUMP_IF_FALSE` 等。VM 实现在 `bytecode.py` 中。'
            ),
            'body_en': (
                'Common instructions: `LOAD_CONST`, `CALL_FUNCTION`, `RETURN_VALUE`, `JUMP_IF_FALSE`, etc. See `bytecode.py` for the VM implementation.'
            ),
            'code': None,
        },
        {
            'title_cn': '编译器与自举',
            'title_en': 'Compiler and Self-hosting (Bootstrapping)',
            'body_cn': (
                '自举流程：在 `bootstrap/` 用 H# 实现 tokenizer/parser/compiler，通过 Python 桥接逐步替换 Python 端编译器。'
            ),
            'body_en': (
                'Bootstrapping flow: implement tokenizer/parser/compiler in H# under `bootstrap/` and use a Python bridge to run and validate; gradually replace the Python compiler.'
            ),
            'code': None,
        },
        {
            'title_cn': '示例与练习',
            'title_en': 'Examples and Exercises',
            'body_cn': (
                '练习：实现 factorial、map/filter 示例、闭包计数器、扩展 bootstrap 编译器以支持新 AST 节点。'
            ),
            'body_en': (
                'Exercises: implement factorial, map/filter examples, closure-based counters, and extend the bootstrap compiler to support new AST nodes.'
            ),
            'code': None,
        },
        {
            'title_cn': '附录：源码引用',
            'title_en': 'Appendix: Source References',
            'body_cn': (
                '源码入口：`lexer.py`, `parser.py`, `compiler.py`, `bytecode.py`, `interpreter.py`。引导实现位于 `bootstrap/`。'
            ),
            'body_en': (
                'Source entry points: `lexer.py`, `parser.py`, `compiler.py`, `bytecode.py`, `interpreter.py`. Bootstrap implementations are in `bootstrap/`.'
            ),
            'code': None,
        },
    ]

    for s in sections:
        doc.add_heading(f"{s['title_cn']} / {s['title_en']}", level=1)
        add_paragraph_cn_en(doc, s['body_cn'], s['body_en'])
        if s.get('code'):
            add_code_block(doc, s['code'])

    # Save
    os.makedirs(os.path.dirname(outpath), exist_ok=True)
    doc.save(outpath)


if __name__ == '__main__':
    out = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs', 'HSharp_Tutorial_bilingual.docx')
    print(f'Generating {out} ...')
    build_doc(out)
    print('Done.')
